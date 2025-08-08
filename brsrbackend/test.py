from fastapi import FastAPI, UploadFile, File, Form, HTTPException,Depends,APIRouter, Body
from fastapi.responses import JSONResponse,FileResponse
from google.api_core.exceptions import ResourceExhausted
from fastapi.middleware.cors import CORSMiddleware
# from Extract_keys import KEYS_TO_EXTRACT_A,KEYS_TO_EXTRACT_B,KEYS_TO_EXTRACT_C1
import Extract_keys
from dotenv import load_dotenv, dotenv_values 
from typing import Dict,List,Optional,Any,Union
from database import engine, SessionLocal
from sqlalchemy import Column, Integer, String, ForeignKey, Text,text,inspect,desc
from sqlalchemy.orm import declarative_base, relationship,Session
from models import Base,create_pdf_model,TableRegistry  # Import if not already done
from pydantic import BaseModel
from schemas import PDFTableUpdate
import json
import pdfplumber
import docx
import google.generativeai as genai
import uvicorn
import os
import time
import pdf
import table_1
import table_2
import table_3
import sec_a
import sec_b
import sec_c
import sys
import asyncio
# from test import app

# if sys.platform == "win32":
#     asyncio.set_eve




class ReportItem(BaseModel):
    name: str
    created_date: Any  # or datetime if created_at is a datetime object
    period: str
    progress: int
    status: str
    section: str

class ReportListResponse(BaseModel):
    reports: List[ReportItem]

class SubmitRequest(BaseModel):
    texts: Union[Dict[str, Any], list]  # Accepts a dictionary or a list
    sectionfind: str
    brsrfilename:Any



load_dotenv()

genai.configure(api_key=os.getenv("CREATE_FIRST_PROJECT"))


app = FastAPI()
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


app = FastAPI(title="Document Extractor API" )
app.add_middleware(
    CORSMiddleware,
    # allow_origins=["http://localhost:3000"],
    allow_origins=["http://192.168.2.118:3000","http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_json_from_text(text):
    brace_stack = []
    start_index = None
    for i, char in enumerate(text):
        if char == '{':
            if start_index is None:
                start_index = i
            brace_stack.append('{')
        elif char == '}':
            if brace_stack:
                brace_stack.pop()
                if not brace_stack:
                    json_str = text[start_index:i + 1]
                    try:
                        return json.loads(json_str)
                    except json.JSONDecodeError:
                        continue
    return None

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])


def chunk_text(text,max_tokens=3000):
    paragraphs = text.split("\n" )
    chunks, current_chunk = [], ""
    for para in paragraphs:
        if len(current_chunk) + len(para) < max_tokens:
            current_chunk += para + "\n"
        else:
            chunks.append(current_chunk)
            # print("else part",len(current_chunk),current_chunk)
            current_chunk = para + "\n"
    if current_chunk:
        chunks.append(current_chunk)
    return chunks




def extract_fields_with_gemini_a(text_chunk: str) -> dict:
    # print(text_chunk)
    model = genai.GenerativeModel("gemini-2.0-flash" )
    prompt = f"""
You are an expert in information extraction. Extract the following details from the provided text and return them in valid JSON format with keys exactly as listed below. Only return the JSON — no extra commentary.

{Extract_keys.KEYS_TO_EXTRACT_A}

TEXT:
{text_chunk}
"""
    try:
        response = model.generate_content(prompt)
        # print(response.text)
        parsed = extract_json_from_text(response.text)
        return parsed if parsed else {}
    except ResourceExhausted:
        raise HTTPException(status_code=429, detail="Gemini API quota exceeded." )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



def extract_fields_with_gemini_b(text_chunk: str) -> dict:
    # print(text_chunk)
    model = genai.GenerativeModel("gemini-2.0-flash" )
    prompt = f"""
You are an expert in information extraction. Extract the following details from the provided text and return them in valid JSON format with keys exactly as listed below. Only return the JSON — no extra commentary.

{Extract_keys.KEYS_TO_EXTRACT_B}

TEXT:
{text_chunk}
"""
    try:
        response = model.generate_content(prompt)
        # print(response.text)
        parsed = extract_json_from_text(response.text)
        return parsed if parsed else {}
    except ResourceExhausted:
        raise HTTPException(status_code=429, detail="Gemini API quota exceeded." )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



def extract_fields_with_gemini_c(text_chunk: str,principle_key:str) -> dict:
    print("This is key",principle_key)
    def extract_key_fun_c(principle_key):
      extract_key_fun_c={
      "principle_1":Extract_keys.KEYS_TO_EXTRACT_C1,
      "principle_2":Extract_keys.KEYS_TO_EXTRACT_C2,
      "principle_3":Extract_keys.KEYS_TO_EXTRACT_C3,
      "principle_4":Extract_keys.KEYS_TO_EXTRACT_C4,
      "principle_5":Extract_keys.KEYS_TO_EXTRACT_C5,
      "principle_6":Extract_keys.KEYS_TO_EXTRACT_C6,
      "principle_7":Extract_keys.KEYS_TO_EXTRACT_C7,
      "principle_8":Extract_keys.KEYS_TO_EXTRACT_C8,
      "principle_9":Extract_keys.KEYS_TO_EXTRACT_C9,

      }
      return extract_key_fun_c[principle_key]
    
    extract_section_key=extract_key_fun_c(principle_key)
    
    model = genai.GenerativeModel("gemini-2.0-flash" )
    prompt = f"""
You are an expert in information extraction. Extract the following details from the provided text and return them in valid JSON format with keys exactly as listed below. Only return the JSON — no extra commentary.

{extract_section_key}

TEXT:
{text_chunk}

"""
    try:
        response = model.generate_content(prompt)
        # print(response.text)
        parsed = extract_json_from_text(response.text)
        return parsed if parsed else {}
    except ResourceExhausted:
        raise HTTPException(status_code=429, detail="Gemini API quota exceeded." )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



def merge_results(results):
    final = {}
    for result in results:
        for k, v in result.items():
            if k not in final or not final[k]:
                final[k] = v
    return final


    
def parse_brsr_text_section_a(file_path,json_merge):
  print("sucessfully file sent",file_path,type(file_path))
  result=table_1.llama_parse_function(file_path)
  print("##########",result)
  return {
    "data": 
      {
        "title": "GENERAL DISCLOSURES",
        "section": "section_a",
        "parts": [
          {
            "partNo": "one",
            "subtitle": "Details of the listed entity",
            "questions": [
              {
                "questionNo": "1",
                "question": "Corporate Identity Number (CIN) of the Listed Entity",
                "questionAnswer":  json_merge["Corporate Identity Number (CIN) of the Listed Entity"]
              },
              {
                "questionNo": "2",
                "question": "Name of the Listed Entity",
                "questionAnswer":json_merge["Name of the Listed Entity"]
              },
              {
                "questionNo": "3",
                "question": "Year of incorporation",
                "questionAnswer":json_merge["Year of incorporation"]
              },
              {
                "questionNo": "4",
                "question": "Registered office address",
                "questionAnswer":json_merge["Registered office address"]
              },
              {
                "questionNo": "5",
                "question": "Corporate address",
                "questionAnswer":json_merge["Corporate address"]
              },
              {
                "questionNo": "6",
                "question": "Email",
                "questionAnswer":json_merge["Email"]
              },
              {
                "questionNo": "7",
                "question": "Telephone",
                "questionAnswer":json_merge["Telephone"]
              },
              {
                "questionNo": "8",
                "question": "Website",
                "questionAnswer":json_merge["Website"]
              },
              {
                "questionNo": "9",
                "question": "Financial year for which reporting is being done",
                "questionAnswer":json_merge["Financial year for which reporting is being done"]
              },
              {
                "questionNo": "10",
                "question": "Name of the Stock Exchange(s) where shares are listed",
                "questionAnswer":json_merge["Name of the Stock Exchange(s) where shares are listed"]
              },
              {
                "questionNo": "11",
                "question": "Paid-up Capital",
                "questionAnswer":json_merge["Paid-up Capital"]
              },
              {
                "questionNo": "12",
                "question": "Name and contact details of the person who may be contacted",
                "questionAnswer":json_merge["Name and contact details"]
              },
              {
                "questionNo": "13",
                "question": "Reporting boundary - Are the disclosures under this report made on a standalone basis (i.e. only for the entity) or on a consolidated basis (i.e. for the entity and all the entities which form a part of its consolidated financial statements, taken together)",
                "questionAnswer":json_merge["Reporting boundary - Are the disclosures under this report made on a standalone basis (i.e. only for the entity) or on a consolidated basis (i.e. for the entity and all the entities which form a part of its consolidated financial statements, taken together)"]
              },
              {
                "questionNo": "14",
                "question": "Name of assurance provider",
                "questionAnswer":json_merge["Name of assurance provider"]
              },
              {
                "questionNo": "15",
                "question": "Type of assurance obtained",
                "questionAnswer":json_merge["Type of assurance obtained"]
              }
            ]
          },

          {
            "partNo": "two",
            "subtitle": "Products / Services",
            "questions": [
              {
                "questionNo": "1",
                "question":"Details of business activities (accounting for 90% of the turnover):",
                "questionAnswer":table_1.table(result,sec_a.Details_of_business),

              },
              {
                "questionNo": "2",
                "question": "Products/Services sold by the entity (accounting for 90% of the entity’s Turnover):",
                "questionAnswer":table_1.table(result,sec_a.Products_Services),

              }
            ]
          },
          {
            "partNo": "three",
            "subtitle": "Operations",
            "questions": [
              {
                "questionNo": "1",
                "question": "Number of locations where plants and/or operations/offices of the entity are situated",
                "questionAnswer":table_1.table(result,sec_a.Number_of_locations_where),
              },
              {
                "questionNo": "2",
                "question": "Number of locations",
                "questionAnswer":table_1.table(result,sec_a.Number_of_locations),
              },
              {
                "questionNo": "3",
                "question": "What is the contribution of exports as a percentage of the total turnover of the entity?",
                "questionAnswer":json_merge["What is the contribution of exports as a percentage of the total turnover of the entity?"],
              },
              {
                "questionNo": "4",
                "question":"A brief on types of customers",
                "questionAnswer":json_merge["A brief on types of customers"],
              },
            ]
          },
          {
            "partNo": "four",
            "subtitle": "Employees",
            "questions": [
              {
              "questionNo": "1",
              "question": "Employees and workers (including differently abled):",
              "questionAnswer":table_1.table(result,sec_a.Employees_and_workers),
              },
              {
              "questionNo": "2",
              "question": "Differently abled Employees and workers:",
              "questionAnswer":table_1.table(result,sec_a.Differently_abled_Employees),
              },
              {
                "questionNo": "3",
                "question": "Participation/Inclusion/Representation of women",
                "questionAnswer":table_1.table(result,sec_a.Participation_Inclusion),
              },
              {
                "questionNo": "4",
                "question": "Turnover rate for permanent employees and workers (Disclose trends for the past 3 years)",
                "questionAnswer":table_1.table(result,sec_a.Turnover_rate_for_permanent),
              }
            ]
          },
          {
            "partNo": "five",
            "subtitle": "Holding, Subsidiary and Associate Companies (including joint ventures)",
            "questions": [
              {
                "questionNo": "1",
                "question": "Names of holding/subsidiary/associate companies/joint ventures",
                "questionAnswer":table_1.table(result,sec_a.Names_of_holding),
              }
            ]
          },
          {
            "partNo": "six",
            "subtitle": "CSR Details",
            "questions": [
            {
            "questionNo": "1",
            "question": "Whether CSR is applicable as per section 135 of Companies Act, 2013: (Yes/No)",
            "questionAnswer":json_merge["Whether CSR is applicable as per section 135 of Companies Act, 2013: (Yes/No)"],
            },
                        {
            "questionNo": "2",
            "question": "Turnover (in Rs.)",
            "questionAnswer":json_merge["Turnover (in Rs.)"],
            },
            {
            "questionNo": "3",
            "question": "Net worth (in Rs.)",
            "questionAnswer":json_merge["Net worth (in Rs.)"],
            },
            ]
          },
          {
            "partNo": "seven",
            "subtitle": "Transparency and Disclosures Compliances",
            "questions": [
              {
                "questionNo": "1",
                "question": "Complaints/Grievances on any of the principles (Principles 1 to 9) under the National Guidelines on Responsible Business Conduct:",
                "questionAnswer":table_1.table(result,sec_a.Complaints_Grievances),
              },
              {
            "questionNo": "2",
            "question": "Please indicate material responsible business conduct and sustainability issues pertaining to environmental and social matters that present a risk or an opportunity to your business, rationale for identifying the same, approach to adapt or mitigate the risk along-with its financial implications, as per the following format.",
            "questionAnswer":table_1.table(result,sec_a.Please_indicate_material),
            },

            ]
          }
        ]
      },
  
 }


def parse_brsr_text_section_b(file_path,json_merge):
    print("sucessfully file sent",file_path)
    result=table_2.llama_parse_function(file_path)
    print("##################",result)


    return{
      "data":
            {
        "title": "MANAGEMENT AND PROCESS DISCLOSURES",
        "section": "section_b",
        "parts": [
          {
            "partNo": "one",
            "subtitle": "Policy and management processes",
            "questions": [
              {
                "questionNo": "1",
                "question": "Whether your entity’s policy/policies cover each principle and its core elements of the NGRBCs. (Yes/No)",
                
                "questionAnswer":result["Whether your entity’s policy/policies cover each principle and its core elements of the NGRBCs"]
              },
              {
                "questionNo": "2",
                "question": "Has the policy been approved by the Board? (Yes/No)",
                
                "questionAnswer":result["Has the policy been approved by the Board?"]
              },
              {
                "questionNo": "3",
                "question": "Web Link of the Policies, if available",
                
                "questionAnswer":json_merge["Web Link of the Policies, if available"],
              },
              {
                "questionNo": "4",
                "question": "Whether the entity has translated the policy into procedures. (Yes / No)",
                
                "questionAnswer":result["Whether the entity has translated the policy into procedures"],
              },
              {
                "questionNo": "5",
                "question": "Do the enlisted policies extend to your value chain partners? (Yes/No)",
                
                "questionAnswer":result["Do the enlisted policies extend to your value chain partners?"]
              },
              {
                "questionNo": "6",
                "question": "Name of the national and international codes/ certifications/labels/ standards (e.g. Forest Stewardship Council, Fairtrade, Rainforest Alliance,Trustea) standards (e.g. SA 8000, OHSAS, ISO, BIS) adopted by your entity and mapped to each principle.",
                
                "questionAnswer":result["Name of the national and international codes/certifications/labels/standards adopted by your entity and mapped to each principle"]
              },
              {
                "questionNo": "7",
                "question": "Specific commitments, goals and targets set by the entity with defined timelines, if any.",
                
                "questionAnswer":result["Specific commitments, goals and targets set by the entity with defined timelines, if any"]
              },
              {
                "questionNo": "8",
                "question": "Performance of the entity against the specific commitments, goals and targets along-with reasons in case the same are not met.",
                
                "questionAnswer":json_merge["Performance of the entity against the specific commitments, goals and targets along-with reasons in case the same are not met."],
              },

            ]
          },
          {
            "partNo":"two",
            "subtitle": "Governance, leadership and oversight",
            "questions": [
              {
                "questionNo": "1",
                "question": "Statement by director responsible for the business responsibility report, highlighting ESG related challenges, targets and achievements (listed entity has flexibility regarding the placement of this disclosure)",
                
                "questionAnswer":json_merge["Statement by director responsible for the business responsibility report, highlighting ESG related challenges, targets and achievements"],
              },
              {
                "questionNo": "2",
                "question": "Details of the highest authority responsible for implementation and oversight of the Business Responsibility policy (ies).",
                
                "questionAnswer":json_merge["Details of the highest authority responsible for implementation and oversight of the Business Responsibility policy (ies)."]
              },
              {
                "questionNo": "3",
                "question": "Does the entity have a specified Committee of the Board/ Director responsible for decision making on sustainability related issues? (Yes / No). If yes, provide details.",
                
                "questionAnswer":result["Does the entity have a specified Committee of the Board/ Director responsible for decision making on sustainability related issues?"]
              },
              {
                "questionNo": "4",
                "question": "Indicate whether review was undertaken by Director / Committee of the Board/ Any other Committee",
                "questionAnswer":result["Indicate whether review was undertaken by Director / Committee of the Board/ Any other Committee"],
              },
              {
                "questionNo": "5",
                "question": "Frequency(Annually/ Half yearly/ Quarterly/ Any other – please specify)",
                "questionAnswer":result["Frequency(Annually/ Half yearly/ Quarterly/ Any other – please specify)"]
              },
              {
                "questionNo": "6",
                "question": "Has the entity carried out independent assessment/ evaluation of the working of its policies by an external agency? (Yes/No). If yes, provide name of the agency.",
                
                "questionAnswer":result["Has the entity carried out independent assessment/ evaluation of the working of its policies by an external agency? (Yes/No). If yes, provide name of the agency"]
              },
              {
                "questionNo": "7",
                "question": "If answer to question (1) above is “No” i.e. not all Principles are covered by a policy, reasons to be stated, as below:",
                "questionAnswer":result["If answer to question (1) above is “No” i.e. not all Principles are covered by a policy, reasons to be stated, as below"]
              },
              {
                "questionNo": "8",
                "question": "Upstream (Suppliers & Logistics Partners)",                
                "questionAnswer":json_merge["Upstream (Suppliers & Logistics Partners)"],
              },
              {
                "questionNo": "9",
                "question": "Downstream (Distributors & Customers)",                
                "questionAnswer":json_merge["Downstream (Distributors & Customers)"],
              },
              ]},]}}   

def parse_brsr_text_section_c1(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,1)
  print("#####################",result)
  
  
  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
            {
            "partNo": "one",
            "subtitle": "Businesses should conduct and govern themselves with integrity, and in a manner that is Ethical, Transparent and Accountable",
            "questions": [
              {
                "questionNo": "1",
                "question": "Percentage coverage by training and awareness programmes on any of the Principles during the financial year:",
                "questionAnswer":table_3.table(result,sec_c.percentage_coverage),#principle_1.Percentage_coverage_by_training(file_path),
              },
              {
              "questionNo":"2",
              "question":"Details of fines / penalties /punishment/ award/ compounding fees/ settlement amount paid in proceedings (by the entity or by directors / KMPs) with regulators/ law enforcement agencies/ judicial institutions, in the financial year, in the following format (Note: the entity shall make disclosures on the basis of materiality as specified in Regulation 30 of SEBI (Listing Obligations and Disclosure Obligations) Regulations, 2015 and as disclosed on the entity’s website",
              "questionAnswer":"",
              },
              {
                "questionNo": "3",
                "question":"Monetary",
                "questionAnswer":table_3.table(result,sec_c.Monetary),
              },
              {
                "questionNo": "4",
                "question":"Non-Monetary",
                "questionAnswer":table_3.table(result,sec_c.Non_Monetary),
              },
              {
                "questionNo": "5",
                "question": "Of the instances disclosed in Question 2 above, details of the Appeal/ Revision preferred in cases where monetary or non-monetary action has been appealed.",
                "questionAnswer":table_3.table(result,sec_c.Of_the_instances_disclosed),
              },
                {
                "questionNo": "6",
                "question": "Does the entity have an anti-corruption or anti-bribery policy? If yes, provide details in brief and if available, provide a web-link to the policy.",
                "questionAnswer":json_merge["Does the entity have an anti-corruption or anti-bribery policy? If yes, provide details in brief and if available, provide a web-link to the policy."],
              },
                {
                "questionNo": "7",
                "question": "Number of Directors/KMPs/employees/workers against whom disciplinary action was taken by any law enforcement agency for the charges of bribery/ corruption",
                "questionAnswer":table_3.table(result,sec_c.Number_of_Directors),
              },
                {
                "questionNo": "8",
                "question": "Details of complaints with regard to conflict of interest",
                "questionAnswer":table_3.table(result,sec_c.Number_of_Complaints),
              },
                {
                "questionNo": "9",
                "question": "Provide details of any corrective action taken or underway on issues related to fines / penalties / action taken by regulators/ law enforcement agencies/ judicial institutions, on cases of corruption and conflicts of interest.",
                "questionAnswer":json_merge["Provide details of any corrective action taken or underway on issues related to fines / penalties / action taken by regulators/ law enforcement agencies/ judicial institutions, on cases of corruption and conflicts of interest."],
              },
                {
                "questionNo": "10",
                "question": "Number of days of accounts payables ((Accounts payable *365) / Cost of goods/services procured) in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Number_of_days),#principle_1.Number_of_days_of_accounts(file_path),
              },
                {
                "questionNo": "11",
                "question": "Provide details of concentration of purchases and sales with trading houses, dealers, and related parties along-with loans and advances & investments, with related parties, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Provide_details_of_concentration),#principle_1.Provide_details_of_concentration_of_purchases(file_path),
              },
                {
                "questionNo": "12",
                "question": "Awareness programmes conducted for value chain partners on any of the Principles during the financial year:",
                "questionAnswer":table_3.table(result,sec_c.Awareness_programmes),#principle_1.Awareness_programmes_conducted(file_path),
              },
                {
                "questionNo": "13",
                "question": "Does the entity have processes in place to avoid/ manage conflict of interests involving members of the Board? (Yes/No) If Yes, provide details of the same.",
                "questionAnswer":json_merge["Does the entity have processes in place to avoid/ manage conflict of interests involving members of the Board? (Yes/No) If Yes, provide details of the same."],#principle_1.Does_the_entity_have_processes(file_path),
              }

            ]
          },
                ]
      }}

def parse_brsr_text_section_c2(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,2)
  print("#####################",result)
  
  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
                   {
            "partNo": "two",
            "subtitle": "Businesses should provide goods and services in a manner that is sustainable and safe",
            "questions": [
              {
                "questionNo": "1",
                "question": "Percentage of R&D and capital expenditure (capex) investments in specific technologies to improve the environmental and social impacts of product and processes to total R&D and capex investments made by the entity, respectively.",
                "questionAnswer":table_3.table(result,sec_c.Percentage_of_RD),
              },
              {
                "questionNo": "2",
                "question": "Does the entity have procedures in place for sustainable sourcing? (Yes/No)",
                "questionAnswer":json_merge["Does the entity have procedures in place for sustainable sourcing? (Yes/No)"]
              },
              {
                "questionNo": "3",
                "question": "If yes, what percentage of inputs were sourced sustainably?",
                "questionAnswer":json_merge["If yes, what percentage of inputs were sourced sustainably?"]
              },
              {
                "questionNo": "4",
                "question": "Describe the processes in place to safely reclaim your products for reusing, recycling and disposing at the end of life, for :",
                "questionAnswer":json_merge["Describe the processes in place to safely reclaim your products for reusing, recycling and disposing at the end of life, for :"],
              },

              {
                "questionNo": "5",
                "question": "Whether Extended Producer Responsibility (EPR) is applicable to the entity’s activities (Yes / No). If yes, whether the waste collection plan is in line with the Extended Producer Responsibility (EPR) plan submitted to Pollution Control Boards? If not, provide steps taken to address the same.",
                "questionAnswer":json_merge["Whether Extended Producer Responsibility (EPR) is applicable to the entity’s activities (Yes / No). If yes, whether the waste collection plan is in line with the Extended Producer Responsibility (EPR) plan submitted to Pollution Control Boards? If not, provide steps taken to address the same."]
              },

              {
                "questionNo": "6",
                "question": "Has the entity conducted Life Cycle Perspective / Assessments (LCA) for any of its products (for manufacturing industry) or for its services (for service industry)? If yes, provide details in the following format?",
                "questionAnswer":table_3.table(result,sec_c.Has_the_entity),
              },

              {
                "questionNo": "7",
                "question": "If there are any significant social or environmental concerns and/or risks arising from production or disposal of your products / services, as identified in the Life Cycle Perspective / Assessments (LCA) or through any other means, briefly describe the same along-with action taken to mitigate the same.",
                "questionAnswer":table_3.table(result,sec_c.If_there_are_any),#json_merge["If there are any significant social or environmental concerns and/or risks arising from production or disposal of your products / services, as identified in the Life Cycle Perspective / Assessments (LCA) or through any other means, briefly describe the same along-with action taken to mitigate the same."],
              },

              {
                "questionNo": "8",
                "question": "Percentage of recycled or reused input material to total material (by value) used in production (for manufacturing industry) or providing services (for service industry)",
                "questionAnswer":table_3.table(result,sec_c.Percentage_of_recycled),#json_merge["Percentage of recycled or reused input material to total material (by value) used in production (for manufacturing industry) or providing services (for service industry)"],
              },

              {
                "questionNo": "9",
                "question": "Of the products and packaging reclaimed at end of life of products, amount (in metric tonnes) reused, recycled, and safely disposed, as per the following format:",
                "questionAnswer":table_3.table(result,sec_c.Of_the_products),#json_merge["Of the products and packaging reclaimed at end of life of products, amount (in metric tonnes) reused, recycled, and safely disposed, as per the following format:"],
              },
              {
                "questionNo": "10",
                "question": "Reclaimed products and their packaging materials (as percentage of products sold) for each product category.",
                "questionAnswer": table_3.table(result,sec_c.Reclaimed_products),#json_merge["Reclaimed products and their packaging materials (as percentage of products sold) for each product category."],
              },
        ]},]}}

def parse_brsr_text_section_c3(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,3)
  print("#####################",result)

  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
           {
            "partNo": "three",
            "subtitle": "Businesses should respect and promote the well-being of all employees, including those in their value chains",
            "questions": [
              {
                "questionNo": "1",
                "question":"Details of measures for the well-being of employees:",
                "questionAnswer":table_3.table(result,sec_c.well_being_of_employees),#json_merge["Details of retirement benefits, for Current and Previous FY"],
              },
              {
                "questionNo": "2",
                "question":"Details of measures for the well-being of workers:",
                "questionAnswer":table_3.table(result,sec_c.well_being_of_workers),
              },
              {
                "questionNo": "3",
                "question": "Details of retirement benefits, for Current and Previous FY",
                "questionAnswer":table_3.table(result,sec_c.Details_of_retirement_benefits),
              },
              {
                "questionNo": "4",
                "question": "Accessibility of workplaces",
                "questionAnswer":json_merge["Accessibility of workplaces"]#principle_3.Accessibility_of_workplaces(file_path),
              },
              {
                "questionNo": "5",
                "question": "Does the entity have an equal opportunity policy as per the Rights of Persons with Disabilities Act, 2016? If so, provide a web-link to the policy.",
                "questionAnswer":json_merge["Does the entity have an equal opportunity policy as per the Rights of Persons with Disabilities Act, 2016? If so, provide a web-link to the policy."],
              },
              {
                "questionNo": "6",
                "question": "Return to work and Retention rates of permanent employees and workers that took parental leave.",
                "questionAnswer":table_3.table(result,sec_c.Return_to_work),
              },
              {
                "questionNo": "7",
                "question": "Is there a mechanism available to receive and redress grievances for the following categories of employees and worker? If yes, give details of the mechanism in brief",
                "questionAnswer":table_3.table(result,sec_c.Is_there_mechanism),#json_merge["Is there a mechanism available to receive and redress grievances for the following categories of employees and worker? If yes, give details of the mechanism in brief"],
              },
              {
                "questionNo": "8",
                "question": "Membership of employees and worker in association(s) or Unions recognised by the listed entity:",
                "questionAnswer":table_3.table(result,sec_c.Membership_employees),#json_merge["Membership of employees and worker in association(s) or Unions recognised by the listed entity:"],
              },
              {
                "questionNo": "9",
                "question": "Details of training given to employees and workers:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_training),#json_merge["Details of training given to employees and workers:"],
              },
              {
                "questionNo": "10",
                "question": "Details of performance and career development reviews of employees and workers:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_performance),#json_merge["Details of performance and career development reviews of employees and workers:"],
              },
              {
                "questionNo": "11",
                "question":"Whether an occupational health and safety management system has been implemented by the entity? (Yes/ No). If yes, the coverage such system?",
                "questionAnswer":json_merge["Whether an occupational health and safety management system has been implemented by the entity? (Yes/ No). If yes, the coverage such system?"]#principle_3.a_Whether_an_occupational(file_path),#json_merge["Details of performance and career development reviews of employees and workers:"],
              },
              {
                "questionNo": "12",
                "question":"What are the processes used to identify work-related hazards and assess risks on a routine and non-routine basis by the entity?",
                "questionAnswer":json_merge["What are the processes used to identify work-related hazards and assess risks on a routine and non-routine basis by the entity?"]#principle_3.b_What_are_the_processes(file_path),#json_merge["Details of performance and career development reviews of employees and workers:"],
              },
              {
                "questionNo": "13",
                "question":"Whether you have processes for workers to report the work related hazards and to remove themselves from such risks. (Y/N)",
                "questionAnswer":json_merge["Whether you have processes for workers to report the work related hazards and to remove themselves from such risks. (Y/N)"]
              },
              {
                "questionNo": "14",
                "question":"Do the employees/ worker of the entity have access to non-occupational medical and healthcare services? (Yes/ No)",
                "questionAnswer":json_merge["Do the employees/ worker of the entity have access to non-occupational medical and healthcare services? (Yes/ No)"]
              },
              {
                "questionNo": "15",
                "question": "Details of safety related incidents, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_safety),
              },
              {
                "questionNo": "16",
                "question": "Describe the measures taken by the entity to ensure a safe and healthy work place",
                "questionAnswer":json_merge["Describe the measures taken by the entity to ensure a safe and healthy work place"],
              },
              {
                "questionNo": "17",
                "question": "Number of Complaints on the following made by employees and workers:",
                "questionAnswer":table_3.table(result,sec_c.Number_of_Complaints),
              },
              {
                "questionNo": "18",
                "question": "Assessments for the year:",
                "questionAnswer":table_3.table(result,sec_c.Assessments_for_the_year),
              },
              {
                "questionNo": "19",
                "question": "Provide details of any corrective action taken or underway to address safety-related incidents (if any) and on significant risks / concerns arising from assessments of health & safety practices and working conditions.",
                "questionAnswer":json_merge["Provide details of any corrective action taken or underway to address safety-related incidents (if any) and on significant risks / concerns arising from assessments of health & safety practices and working conditions."]#principle_3.Provide_details_of_any_corrective(file_path),
              },
                {
                "questionNo": "20",
                "question": "Does the entity extend any life insurance or any compensatory package in the event of death of (A) Employees (Y/N)",
                "questionAnswer":json_merge["Does the entity extend any life insurance or any compensatory package in the event of death of (A) Employees (Y/N)"]
              },
              {
                "questionNo": "21",
                "question": "Provide the measures undertaken by the entity to ensure that statutory dues have been deducted and deposited by the value chain partners",
                "questionAnswer":json_merge["Provide the measures undertaken by the entity to ensure that statutory dues have been deducted and deposited by the value chain partners"]
              },
              {
                "questionNo": "22",
                "question": "Provide the number of employees / workers having suffered high consequence work-related injury / ill-health / fatalities (as reported in Q11 of Essential Indicators above), who have been are rehabilitated and placed in suitable employment or whose family members have been placed in suitable employment:",
                "questionAnswer":table_3.table(result,sec_c.Provide_the_number_employees),
              },
              {
                "questionNo": "23",
                "question": "Does the entity provide transition assistance programs to facilitate continued employability and the management of career endings resulting from retirement or termination of employment? (Yes/ No)",
                "questionAnswer":json_merge["Does the entity provide transition assistance programs to facilitate continued employability and the management of career endings resulting from retirement or termination of employment? (Yes/ No)"]
              },
              {
                "questionNo": "24",
                "question": "Details on assessment of value chain partners:",
                "questionAnswer":table_3.table(result,sec_c.Details_on_assessment),
              },
              {
                "questionNo": "25",
                "question": "Provide details of any corrective actions taken or underway to address significant risks / concerns arising from assessments of health and safety practices and working conditions of value chain partners.",
                "questionAnswer":json_merge["Provide details of any corrective actions taken or underway to address significant risks / concerns arising from assessments of health and safety practices and working conditions of value chain partners."]
              },
        ]
      },
                ]
      }}

def parse_brsr_text_section_c4(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,4)
  print("#####################",result)


  
  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
          
         {
            "partNo": "four",
            "subtitle": "Businesses should respect the interests of and be responsive to all its stakeholders",
            "questions": [
              {
                "questionNo": "1",
                "question": "Describe the processes for identifying key stakeholder groups of the entity.",
                "questionAnswer":json_merge["Describe the processes for identifying key stakeholder groups of the entity."]
              },
              {
                "questionNo": "2",
                "question": "List stakeholder groups identified as key for your entity and the frequency of engagement with each stakeholder group.",
                "questionAnswer":table_3.table(result,sec_c.List_stakeholder),
              },
              {
                "questionNo": "3",
                "question": "Provide the processes for consultation between stakeholders and the Board on economic, environmental, and social topics or if consultation is delegated, how is feedback from such consultations provided to the Board.",
                "questionAnswer":json_merge["Provide the processes for consultation between stakeholders and the Board on economic, environmental, and social topics or if consultation is delegated, how is feedback from such consultations provided to the Board."]#principle_4.Provide_the_processes(file_path),#json_merge["Provide the processes for consultation between stakeholders and the Board on economic, environmental, and social topics or if consultation is delegated, how is feedback from such consultations provided to the Board."],
              },
              {
                "questionNo": "4",
                "question": "Whether stakeholder consultation is used to support the identification and management of environmental, and social topics (Yes / No). If so, provide details of instances as to how the inputs received from stakeholders on these topics were incorporated into policies and activities of the entity",
                "questionAnswer":json_merge["Whether stakeholder consultation is used to support the identification and management of environmental, and social topics (Yes / No). If so, provide details of instances as to how the inputs received from stakeholders on these topics were incorporated into policies and activities of the entity"]
              },
              {
                "questionNo": "5",
                "question": "Provide details of instances of engagement with, and actions taken to, address the concerns of vulnerable/ marginalized stakeholder groups.",
                "questionAnswer":json_merge["Provide details of instances of engagement with, and actions taken to, address the concerns of vulnerable/ marginalized stakeholder groups."]
              },


        ]
      },
                ]
      }}

def parse_brsr_text_section_c5(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,5)
  print("#####################",result)

  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
          {
            "partNo": "five",
            "subtitle": " Businesses should respect and promote human rights",
            "questions": [
              {
                "questionNo": "1",
                "question": "Employees and workers who have been provided training on human rights issues and policy(ies) of the entity, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Employees_and_workers),
              },
              {
                "questionNo": "2",
                "question": "Details of minimum wages paid to employees and workers, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_minimum),
              },
              {
                "questionNo": "3",
                "question": "Details of remuneration/salary/wages, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_remuneration),
              },
              {
                "questionNo": "4",
                "question": "Do you have a focal point (Individual/ Committee) responsible for addressing human rights impacts or issues caused or contributed to by the business? (Yes/ No)",
                "questionAnswer":json_merge["Do you have a focal point (Individual/ Committee) responsible for addressing human rights impacts or issues caused or contributed to by the business? (Yes/ No)"]#principle_5.Do_you_have_a_focal_point(file_path),
              },
              {
                "questionNo": "5",
                "question": "Describe the internal mechanisms in place to redress grievances related to human rights issues.",
                "questionAnswer":json_merge["Describe the internal mechanisms in place to redress grievances related to human rights issues."]#principle_5.Describe_the_internal_mechanisms(file_path),
              },
              {
                "questionNo": "6",
                "question": "Number of Complaints on the following made by employees and workers:",
                "questionAnswer":table_3.table(result,sec_c.Number_of_Complaints),
              },
              {
                "questionNo": "7",
                "question":"Complaints filed under the Sexual Harassment of Women at Workplace (Prevention, Prohibition and Redressal) Act, 2013, in the following format:",
                "questionAnswer":"KIKIKIKIKI",
              },
              {
                "questionNo": "8",
                "question": "Mechanisms to prevent adverse consequences to the complainant in discrimination and harassment cases.",
                "questionAnswer":json_merge["Mechanisms to prevent adverse consequences to the complainant in discrimination and harassment cases."]#principle_5.Mechanisms_prevent(file_path),
              },
              {
                "questionNo": "9",
                "question": "Do human rights requirements form part of your business agreements and contracts? (Yes/ No)",
                "questionAnswer":json_merge["Do human rights requirements form part of your business agreements and contracts? (Yes/ No)"]#principle_5.Do_human_rights_requirements(file_path),
              },
              {
                "questionNo": "10",
                "question": "Assessments for the year:",
                "questionAnswer":table_3.table(result,sec_c.Assessments_for_the_year),
              },
              {
                "questionNo": "11",
                "question": "Provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at Question 10 above.",
                "questionAnswer":json_merge["Provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at Question 10 above."]#principle_5.Provide_details_of_any(file_path),
              },
              {
                "questionNo": "12",
                "question": "Details of a business process being modified / introduced as a result of addressing human rights grievances/complaints.",
                "questionAnswer":json_merge["Details of a business process being modified / introduced as a result of addressing human rights grievances/complaints."]#principle_5.Details_of_a_business_process(file_path),
              },
              {
                "questionNo": "13",
                "question": "Details of the scope and coverage of any Human rights due-diligence conducted",
                "questionAnswer":json_merge["Details of the scope and coverage of any Human rights due-diligence conducted"]#principle_5.Details_of_the_scope(file_path),
              },
              {
                "questionNo": "14",
                "question": "Is the premise/office of the entity accessible to differently abled visitors, as per the requirements of the Rights of Persons with Disabilities Act, 2016?",
                "questionAnswer":json_merge["Is the premise/office of the entity accessible to differently abled visitors, as per the requirements of the Rights of Persons with Disabilities Act, 2016?"]#principle_5.Is_the_premise(file_path),
              },
              {
                "questionNo": "15",
                "question": "Details on assessment of value chain partners:",
                "questionAnswer":table_3.table(result,sec_c.Details_on_assessment),
              },
              {
                "questionNo": "16",
                "question": "Provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at Question 4 above.",
                "questionAnswer":json_merge["Provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at Question 4 above."]#principle_5.Provide_details_of_any_corrective(file_path),
              },

        ]
      },
                ]
      }}

def parse_brsr_text_section_c6(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,6)
  print("#####################",result)

  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
            {
            "partNo": "six",
            "subtitle": "Businesses should respect and make efforts to protect and restore",
            "questions": [
              {
                "questionNo": "1",
                "question": "Details of total energy consumption (in Joules or multiples) and energy intensity, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_total_energy),#json_merge["Details of total energy consumption (in Joules or multiples) and energy intensity, in the following format:"],
              },
              {
                "questionNo": "2",
                "question": "Does the entity have any sites / facilities identified as designated consumers (DCs) under the Performance, Achieve and Trade (PAT) Scheme of the Government of India? (Y/N) If yes, disclose whether targets set under the PAT scheme have been achieved. In case targets have not been achieved, provide the remedial action taken, if any.",
                "questionAnswer":json_merge["Does the entity have any sites / facilities identified as designated consumers (DCs) under the Performance, Achieve and Trade (PAT) Scheme of the Government of India? (Y/N) If yes, disclose whether targets set under the PAT scheme have been achieved. In case targets have not been achieved, provide the remedial action taken, if any."]#principle_6.Does_the_entity_have_any_sites(file_path),#json_merge["Does the entity have any sites / facilities identified as designated consumers (DCs) under the Performance, Achieve and Trade (PAT) Scheme of the Government of India? (Y/N) If yes, disclose whether targets set under the PAT scheme have been achieved. In case targets have not been achieved, provide the remedial action taken, if any."],
              },
              {
                "questionNo": "3",
                "question": "Provide details of the following disclosures related to water, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Provide_details_of_the_following),#json_merge["Provide details of the following disclosures related to water, in the following format:"],
              },

              {
                "questionNo": "4",
                "question": "Provide the following details related to water discharged:",
                "questionAnswer":table_3.table(result,sec_c.Provide_the_following),#json_merge["Provide the following details related to water discharged:"],
              },

              {
                "questionNo": "5",
                "question": "Has the entity implemented a mechanism for Zero Liquid Discharge? If yes, provide details of its coverage and implementation",
                "questionAnswer":json_merge["Has the entity implemented a mechanism for Zero Liquid Discharge? If yes, provide details of its coverage and implementation"]#principle_6.Has_the_entity_implemented(file_path),#json_merge["Has the entity implemented a mechanism for Zero Liquid Discharge? If yes, provide details of its coverage and implementation"],
              },

              {
                "questionNo": "6",
                "question": "Please provide details of air emissions (other than GHG emissions) by the entity, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Please_provide_details),#json_merge["Please provide details of air emissions (other than GHG emissions) by the entity, in the following format:"],
              },

              {
                "questionNo": "7",
                "question": "Provide details of greenhouse gas emissions (Scope 1 and Scope 2 emissions) & its intensity, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Provide_details_of_greenhouse),#json_merge["Provide details of greenhouse gas emissions (Scope 1 and Scope 2 emissions) & its intensity, in the following format:"],
              },

              {
                "questionNo": "8",
                "question": "Does the entity have any project related to reducing Green House Gas emission? If Yes, then provide details.",
                "questionAnswer":json_merge["Does the entity have any project related to reducing Green House Gas emission? If Yes, then provide details."]#principle_6.Does_the_entity_have_any_project(file_path),#json_merge["Does the entity have any project related to reducing Green House Gas emission? If Yes, then provide details."],
              },

              {
                "questionNo": "9",
                "question": "Provide details related to waste management by the entity, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Provide_details_related),#json_merge["Provide details related to waste management by the entity, in the following format:"],
              },

              {
                "questionNo": "10",
                "question": "Briefly describe the waste management practices adopted in your establishments. Describe the strategy adopted by your company to reduce usage of hazardous and toxic chemicals in your products and processes and the practices adopted to manage such wastes.",
                "questionAnswer":json_merge["Briefly describe the waste management practices adopted in your establishments. Describe the strategy adopted by your company to reduce usage of hazardous and toxic chemicals in your products and processes and the practices adopted to manage such wastes."]#principle_6.Briefly_describe_the_waste(file_path),#json_merge["Briefly describe the waste management practices adopted in your establishments. Describe the strategy adopted by your company to reduce usage of hazardous and toxic chemicals in your products and processes and the practices adopted to manage such wastes."],
              },

              {
                "questionNo": "11",
                "question": "If the entity has operations/offices in/around ecologically sensitive areas (such as national parks, wildlife sanctuaries, biosphere reserves, wetlands, biodiversity hotspots, forests, coastal regulation zones etc.) where environmental approvals / clearances are required, please specify details in the following format:",
                "questionAnswer":table_3.table(result,sec_c.If_the_entity_has_operations),
              },

              {
                "questionNo": "12",
                "question": "Details of environmental impact assessments of projects undertaken by the entity based on applicable laws, in the current financial year:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_environmental),#json_merge["Details of environmental impact assessments of projects undertaken by the entity based on applicable laws, in the current financial year:"],
              },

              {
                "questionNo": "13",
                "question": "Is the entity compliant with the applicable environmental law/ regulations/ guidelines in India; such as the Water (Prevention and Control of Pollution) Act, Air (Prevention and Control of Pollution) Act, Environment protection act and rules thereunder (Y/N). If not, provide details of all such non-compliances, in the following format:",
                "questionAnswer":table_3.table(result,sec_c.Is_the_entity_compliant),#json_merge["Is the entity compliant with the applicable environmental law/ regulations/ guidelines in India; such as the Water (Prevention and Control of Pollution) Act, Air (Prevention and Control of Pollution) Act, Environment protection act and rules thereunder (Y/N). If not, provide details of all such non-compliances, in the following format:"],
              },

              {
                "questionNo": "14",
                "question": "Water withdrawal, consumption and discharge in areas of water stress (in kilolitres):",
                "questionAnswer":table_3.table(result,sec_c.Water_withdrawal),#json_merge["Water withdrawal, consumption and discharge in areas of water stress (in kilolitres):"],
              },
              

              {
                "questionNo": "15",
                "question": "Please provide details of total Scope 3 emissions & its intensity, in the following format",
                "questionAnswer":table_3.table(result,sec_c.Please_provide_details),#json_merge["Please provide details of total Scope 3 emissions & its intensity, in the following format"],
              },

              {
                "questionNo": "16",
                "question": "With respect to the ecologically sensitive areas reported at Question 10 of Essential Indicators above, provide details of significant direct & indirect impact of the entity on biodiversity in such areas along-with prevention and remediation activities",
                "questionAnswer":json_merge["With respect to the ecologically sensitive areas reported at Question 10 of Essential Indicators above, provide details of significant direct & indirect impact of the entity on biodiversity in such areas along-with prevention and remediation activities"],
              },

              {
                "questionNo": "17",
                "question": "If the entity has undertaken any specific initiatives or used innovative technology or solutions to improve resource efficiency, or reduce impact due to emissions / effluent discharge / waste generated, please provide details of the same as well as outcome of such initiatives, as per the following format:",
                "questionAnswer":table_3.table(result,sec_c.If_the_entity_has_undertaken),#json_merge["If the entity has undertaken any specific initiatives or used innovative technology or solutions to improve resource efficiency, or reduce impact due to emissions / effluent discharge / waste generated, please provide details of the same as well as outcome of such initiatives, as per the following format:"],
              },

              {
                "questionNo": "18",
                "question": "Does the entity have a business continuity and disaster management plan? Give details in 100 words/ web link.",
                "questionAnswer":json_merge["Does the entity have a business continuity and disaster management plan? Give details in 100 words/ web link."]
              },

              {
                "questionNo": "19",
                "question": "Disclose any significant adverse impact to the environment, arising from the value chain of the entity. What mitigation or adaptation measures have been taken by the entity in this regard.",
                "questionAnswer":json_merge["Disclose any significant adverse impact to the environment, arising from the value chain of the entity. What mitigation or adaptation measures have been taken by the entity in this regard."]
              },
              {
                "questionNo": "20",
                "question": "Percentage of value chain partners (by value of business done with such partners) that were assessed for environmental impacts.",
                "questionAnswer":json_merge["Percentage of value chain partners (by value of business done with such partners) that were assessed for environmental impacts."],
              },
              {
                "questionNo": "21",
                "question": "How many Green Credits have been generated or procured:",
                "questionAnswer":""
              },
              {
                "questionNo": "22",
                "question": "By the listed entity",
                "questionAnswer":""
              },
              {
                "questionNo": "23",
                "question": "By the top ten (in terms of value of purchases and sales,respectively) value chain partners”",
                "questionAnswer":""
              },
        ]
      },
                ]
      }}

def parse_brsr_text_section_c7(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,7)
  print("#####################",result)
  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
            {
            "partNo": "seven",
            "subtitle": " Businesses, when engaging in influencing public and regulatory policy, should do so in a manner that is responsible and transparent",
            "questions": [
              {
                "questionNo": "1",
                "question": "Number of affiliations with trade and industry chambers/ associations.",
                "questionAnswer":json_merge["Number of affiliations with trade and industry chambers/ associations."]#principle_7.Number_of_affiliations(file_path),
              },
              {
                "questionNo": "2",
                "question": "List the top 10 trade and industry chambers/ associations (determined based on the total members of such body) the entity is a member of/ affiliated to, in the following format",
                "questionAnswer":table_3.table(result,sec_c.List_the_top_10_trade),
              },
              
              {
                "questionNo": "3",
                "question": "Provide details of corrective action taken or underway on any issues related to anti-competitive conduct by the entity, based on adverse orders from regulatory authorities.",
                "questionAnswer":table_3.table(result,sec_c.Provide_details_of_corrective_action),
              },
              {
                "questionNo": "4",
                "question": "Details of public policy positions advocated by the entity:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_public_policy),
              },]},]}}

def parse_brsr_text_section_c8(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,8)
  print("#####################",result)

  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
            {
            "partNo": "eight",
            "subtitle": "Businesses should promote inclusive growth and equitable development",
            "questions": [
              {
                "questionNo": "1",
                "question": "Details of Social Impact Assessments (SIA) of projects undertaken by the entity based on applicable laws, in the current financial year.",
                "questionAnswer":table_3.table(result,sec_c.Details_of_Social),
              },
              {
                "questionNo": "2",
                "question": "Provide information on project(s) for which ongoing Rehabilitation and Resettlement (R&R) is being undertaken by your entity, in the following format",
                "questionAnswer":table_3.table(result,sec_c.Provide_information),
                },

              {
                "questionNo": "3",
                "question": "Describe the mechanisms to receive and redress grievances of the community.",
                "questionAnswer":json_merge["Describe the mechanisms to receive and redress grievances of the community."]#principle_8.Describe_the_mechanisms_to_receive(file_path),
                },

              {
                "questionNo": "4",
                "question": "Percentage of input material (inputs to total inputs by value) sourced from suppliers",
                "questionAnswer":table_3.table(result,sec_c.Percentage_of_input),
                },

              {
                "questionNo": "5",
                "question": "Job creation in smaller towns – Disclose wages paid to persons employed (including employees or workers employed on a permanent or non-permanent / on contract basis) in the following locations, as % of total wage cost",
                "questionAnswer":table_3.table(result,sec_c.Job_creation),
                },

              {
                "questionNo": "6",
                "question": "Provide details of actions taken to mitigate any negative social impacts identified in the Social Impact Assessments (Reference: Question 1 of Essential Indicators above):",
                "questionAnswer":table_3.table(result,sec_c.Provide_details_of_actions),
                },

              {
                "questionNo": "7",
                "question": "Provide the following information on CSR projects undertaken by your entity in designated aspirational districts as identified by government bodies",
                "questionAnswer":table_3.table(result,sec_c.Provide_the_following_information),
              },
              {
                "questionNo": "8",
                "question": "Do you have a preferential procurement policy where you give preference to purchase from suppliers comprising marginalized /vulnerable groups? (Yes/No)",
                "questionAnswer":json_merge["Do you have a preferential procurement policy where you give preference to purchase from suppliers comprising marginalized /vulnerable groups? (Yes/No)"],#principle_8.Do_you_have_a_preferential(file_path),
              },

              {
                "questionNo": "9",
                "question": "From which marginalized /vulnerable groups do you procure?",
                "questionAnswer":json_merge["From which marginalized /vulnerable groups do you procure?"]#principle_8.From_which_marginalized(file_path),
              },
              {
                "questionNo": "10",
                "question": "What percentage of total procurement (by value) does it constitute?",
                "questionAnswer":json_merge["What percentage of total procurement (by value) does it constitute?"]#principle_8.What_percentage_of_total_procurement(file_path),
              },
              {
                "questionNo": "11",
                "question": "Details of the benefits derived and shared from the intellectual properties owned or acquired by your entity (in the current financial year), based on traditional knowledge:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_the_benefits),
              },

              {
                "questionNo": "12",
                "question": "Details of corrective actions taken or underway, based on any adverse order in intellectual property related disputes wherein usage of traditional knowledge is involved.",
                "questionAnswer":table_3.table(result,sec_c.Details_of_corrective),
              },

              {
                "questionNo": "13",
                "question": "Details of beneficiaries of CSR Projects:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_the_benefits),
              },
        ]
      },
                ]
      }}

def parse_brsr_text_section_c9(file_path,json_merge):
  print("sucessfully file sent",file_path)
  result=table_3.llama_parse_function(file_path,9)
  print("#####################",result)
  return {
    "data":
      {
        "title": "Principle wise performance disclosure",
        "section": "section_c",
        "parts": [
            {
            "partNo": "nine",
            "subtitle": "Businesses should engage with and provide value to their consumers in a responsible manner",
            "questions": [
              {
                "questionNo": "1",
                "question": "Describe the mechanisms in place to receive and respond to consumer complaints and feedback.",
                "questionAnswer":json_merge["Describe the mechanisms in place to receive and respond to consumer complaints and feedback."]#principle_9.Describe_the_mechanisms(file_path),
                },
              {
                "questionNo": "2",
                "question": "Turnover of products and/ services as a percentage of turnover from all products/service that carry information about:",
                "questionAnswer":table_3.table(result,sec_c.Turnover_of_products),
              },
              {
                "questionNo": "3",
                "question": "Number of consumer complaints in respect of the following:",
                "questionAnswer":table_3.table(result,sec_c.Number_of_consumer),
              },

              {
                "questionNo": "4",
                "question": "Details of instances of product recalls on account of safety issues:",
                "questionAnswer":table_3.table(result,sec_c.Details_of_instances),
              },
              {
                "questionNo": "5",
                "question": "Does the entity have a framework/policy on cyber security and risks related to data privacy? (Yes/No). If available, provide weblink of the policy.",
                "questionAnswer":json_merge["Does the entity have a framework/policy on cyber security and risks related to data privacy? (Yes/No). If available, provide weblink of the policy."]#principle_9.Does_the_entity_have_a_framework(file_path),
              },

              {
                "questionNo": "6",
                "question": "Provide details of any corrective actions taken or underway on issues relating to advertising, and delivery of essential services; cyber security and data privacy of customers; re-occurrence of instances of product recalls; penalty / action taken by regulatory authorities on safety of products / services.",
                "questionAnswer":json_merge["Provide details of any corrective actions taken or underway on issues relating to advertising, and delivery of essential services; cyber security and data privacy of customers; re-occurrence of instances of product recalls; penalty / action taken by regulatory authorities on safety of products / services."]#principle_9.Provide_details_of_any_corrective(file_path),
              },
              {
                "questionNo": "7",
                "question": "Channels / platforms where information on products and services of the entity can be accessed (provide web link, if available).",
                "questionAnswer":json_merge["Channels / platforms where information on products and services of the entity can be accessed (provide web link, if available)."]#principle_9.Channels_platforms(file_path),
              },

              {
                "questionNo": "8",
                "question": "Steps taken to inform and educate consumers about safe and responsible usage of products and/or services.",
                "questionAnswer":json_merge["Steps taken to inform and educate consumers about safe and responsible usage of products and/or services."]#principle_9.Steps_taken(file_path),
              },

              {
                "questionNo": "9",
                "question": "Mechanisms in place to inform consumers of any risk of disruption/discontinuation of essential services.",
                "questionAnswer":json_merge["Mechanisms in place to inform consumers of any risk of disruption/discontinuation of essential services."]#principle_9.Mechanisms_place(file_path),
              },

              {
                "questionNo": "10",
                "question": "Does the entity display product information on the product over and above what is mandated as per local laws? (Yes/No/Not Applicable) If yes, provide details in brief. Did your entity carry out any survey with regard to consumer satisfaction relating to the major products / services of the entity, significant locations of operation of the entity or the entity as a whole? (Yes/No)",
                "questionAnswer":json_merge["Does the entity display product information on the product over and above what is mandated as per local laws? (Yes/No/Not Applicable) If yes, provide details in brief. Did your entity carry out any survey with regard to consumer satisfaction relating to the major products / services of the entity, significant locations of operation of the entity or the entity as a whole? (Yes/No)"]#principle_9.Does_the_entity_display_product_information(file_path),
              },
              {
                "questionNo": "11",
                "question": "Number of instances of data breaches along-with impact",
                "questionAnswer":""#json_merge["Number of instances of data breaches along-with impact"]#principle_9.Number_of_instances(file_path),
              },
              {
                "questionNo": "12",
                "question": "Percentage of data breaches involving personally identifiable information of customers",
                "questionAnswer":""#json_merge["Percentage of data breaches involving personally identifiable information of customers"]#principle_9.Percentage_data_breaches(file_path),
              },
              ]}]}}




@app.post("/extract/" )
async def extract_document(file: UploadFile = File(...),questionKey: str = Form(...),principleKey:str = Form(...)):
    if not (file.filename.endswith(".pdf" ) or file.filename.endswith(".docx" )):
        raise HTTPException(status_code=400, detail="Only PDF or DOCX files are supported." )



    if questionKey=="section_a":
      start_time = time.time()  # ⏱ Start time
      print(questionKey)
      content = await file.read()
      temp_path =f"section_a {file.filename}"
      print(temp_path)
      
      with open(temp_path, "wb" ) as f:
          f.write(content)

      try:
          if file.filename.endswith(".pdf" ):
              text = extract_text_from_pdf(temp_path)
          else:
              text = extract_text_from_docx(temp_path)

            
          chunks = chunk_text(text)
          results = [extract_fields_with_gemini_a(chunk) for chunk in chunks[:5]]
          merged = merge_results(results)          
          res=parse_brsr_text_section_a(temp_path,merged)
          return {"response":res,"brsrfilename":temp_path}  

          
          
      finally:
        end_time = time.time()  # End time
        total_seconds = end_time - start_time
        minutes = int(total_seconds // 60)
        seconds = int(total_seconds % 60)
        print(f"\n⏱ Execution Time: {minutes} minutes, {seconds} seconds")        # if os.path.exists(temp_path):


    elif questionKey=="section_b":
      print(questionKey)
      start_time = time.time()  # ⏱ Start time
      content = await file.read()
      temp_path =f"section_b {file.filename}"
      
      with open(temp_path, "wb" ) as f:
          f.write(content)

      try:
          if file.filename.endswith(".pdf" ):
              text = extract_text_from_pdf(temp_path)
          else:
              text = extract_text_from_docx(temp_path)

                    

          chunks = chunk_text(text)
          results = [extract_fields_with_gemini_b(chunk) for chunk in chunks[:5]]
          merged = merge_results(results)          
          res=parse_brsr_text_section_b(temp_path,merged)
          return {"response":res,"brsrfilename":temp_path}  
        
      finally:
        end_time = time.time()  # End time
        total_seconds = end_time - start_time
        minutes = int(total_seconds // 60)
        seconds = int(total_seconds % 60)
        print(f"\n⏱ Execution Time: {minutes} minutes, {seconds} seconds")        # if os.path.exists(temp_path):


    elif questionKey=="section_c" and principleKey==principleKey:
      print(questionKey)
      print(principleKey)
      print(type(Extract_keys.KEYS_TO_EXTRACT_C1))
      def principlefun(principlestr):
        principles={
          "principle_1":parse_brsr_text_section_c1,
          "principle_2":parse_brsr_text_section_c2,
          "principle_3":parse_brsr_text_section_c3,
          "principle_4":parse_brsr_text_section_c4,
          "principle_5":parse_brsr_text_section_c5,
          "principle_6":parse_brsr_text_section_c6,
          "principle_7":parse_brsr_text_section_c7,
          "principle_8":parse_brsr_text_section_c8,
          "principle_9":parse_brsr_text_section_c9,
          }
        return principles[principlestr]

    
      principle_fun=principlefun(principleKey)         
      start_time = time.time()  # ⏱ Start time
      content = await file.read()
      temp_path = f"section_c {file.filename}"
      
      with open(temp_path, "wb" ) as f:
          f.write(content)

      try:
          if file.filename.endswith(".pdf" ):
              text = extract_text_from_pdf(temp_path)
          else:
              text = extract_text_from_docx(temp_path)
                    
          chunks = chunk_text(text)
          results = []

          results = [extract_fields_with_gemini_c(chunks,principleKey)]
          merged = merge_results(results)
          res=principle_fun(temp_path,merged)
          return {"response":res,"brsrfilename":temp_path}           

      finally:
        end_time = time.time()  # End time
        total_seconds = end_time - start_time
        minutes = int(total_seconds // 60)
        seconds = int(total_seconds % 60)
        print(f"\n⏱ Execution Time: {minutes} minutes, {seconds} seconds")        # if os.path.exists(temp_path):
      
    else:
      return "SECTION NOT FOUND !"
    

@app.post("/submit/" )
async def extract_document(payload:SubmitRequest,db: Session = Depends(get_db)):
    texts=payload.texts
    if texts:
        print("datas####")
        sectionfind=payload.sectionfind
        print("section find",sectionfind)
        texts = {k.lower(): v for k, v in texts.items()}
    
        # print(texts)
        data= {
  "sections":[
    { 
     "title": "GENERAL DISCLOSURES",
     "section": "SECTION A",
     "Categories": [
        { "partRoman":"I",
        "categoryNo": "one",
        "subtitle": "Details of the listed entity",
        "questions": [
          {
            "questionNo": "1",
            "question": "Corporate Identity Number (CIN) of the Listed Entity",
            "questionAnswer":texts.get("corporate identity number (cin) of the listed entity","Not Applicable")
          },
          {
            "questionNo": "2",
            "question": "Name of the Listed Entity",
            "questionAnswer":texts.get("name of the listed entity","Not Applicable")
          },
          {
            "questionNo": "3",
            "question": "Year of incorporation",
            "questionAnswer":texts.get("year of incorporation","Not Applicable")
          },
          {
            "questionNo": "4",
            "question": "Registered office address",
            "questionAnswer":texts.get("registered office address","Not Applicable")
          },
          {
            "questionNo": "5",
            "question": "Corporate address",
            "questionAnswer":texts.get("corporate address","Not Applicable")
            },
          {
            "questionNo": "6",
            "question": "Email",
            "questionAnswer":texts.get("email","Not Applicable")
          },

          {
            "questionNo": "7",
            "question": "Telephone",
            "questionAnswer":texts.get("telephone","Not Applicable")
          },
          {
            "questionNo": "8",
            "question": "Website",
            "questionAnswer":texts.get("website","Not Applicable")
          },
          {
            "questionNo": "9",
            "question": "Financial year for which reporting is being done",
            "questionAnswer":texts.get("financial year for which reporting is being done","Not Applicable")
          },
          {
            "questionNo": "10",
            "question": "Name of the Stock Exchange(s) where shares are listed",
            "questionAnswer":texts.get("name of the stock exchange(s) where shares are listed","Not Applicable")
          },
          {
            "questionNo": "11",
            "question": "Paid-up Capital",
            "questionAnswer":texts.get("paid-up capital","Not Applicable")
          },
          {
            "questionNo": "12",
            "question": "Name and contact details (telephone, email address) of the person who may be contacted in case of any queries on the brsr report",
            "questionAnswer":texts.get("Name and contact details (telephone, email address) of the person who may be contacted in case of any queries on the BRSR report","Not Applicable")
          },
          {
            "questionNo": "13",
            "question": "Reporting boundary - Are the disclosures under this report made on a standalone basis (i.e. only for the entity) or on a consolidated basis (i.e. for the entity and all the entities which form a part of its consolidated financial statements, taken together)",
            "questionAnswer":texts.get("Reporting boundary - Are the disclosures under this report made on a standalone basis (i.e. only for the entity) or on a consolidated basis (i.e. for the entity and all the entities which form a part of its consolidated financial statements, taken together).","Not Applicable")
          },
          {
            "questionNo": "14",
            "question": "Name of assurance provider",
            "questionAnswer":texts.get("name of assurance provider","Not Applicable")
          },
          {
            "questionNo": "15",
            "question": "Type of assurance obtained",
            "questionAnswer":texts.get("type of assurance obtained","Not Applicable")
          }]
      },
          { "partRoman":"II",
            "categoryNo": "two",
            "subtitle": "Products / Services",
            "questions": [
              {
                "questionNo": "1",
                "question":"Details of business activities (accounting for 90% of the turnover):",
                "questionAnswer":texts.get("details of business activities (accounting for 90% of the turnover):","Not Applicable"),

              },
              {
                "questionNo": "2",
                "question": "Products/Services sold by the entity (accounting for 90% of the entity’s Turnover):",
                "questionAnswer":texts.get("products/services sold by the entity (accounting for 90% of the entity’s turnover):","Not Applicable")
              }]
            },
          { "partRoman":"III",
            "categoryNo": "three",
            "subtitle": "Operations",
            "questions": [
              {
                "questionNo": "1",
                "question": "Number of locations where plants and offices of the entity are situated:",
                "questionAnswer":texts.get("number of locations where plants and offices of the entity are situated:","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question": "Number of locations",
                "questionAnswer":texts.get("number of locations","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "What is the contribution of exports as a percentage of the total turnover of the entity?",
                "questionAnswer":texts.get("what is the contribution of exports as a percentage of the total turnover of the entity?","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question":"A brief on types of customers",
                "questionAnswer":texts.get("a brief on types of customers","Not Applicable"),
              },
            ]
          },
          { "partRoman":"IV",
            "categoryNo": "four",
            "subtitle": "Employees",
            "questions": [
              {
              "questionNo": "1",
              "question": "Employees and workers (including differently abled):",
              "questionAnswer":texts.get("employees and workers (including differently abled):","Not Applicable"),
              },
              {
              "questionNo": "2",
              "question": "Differently abled Employees and workers:",
              "questionAnswer":texts.get("differently abled employees and workers:","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "Participation/Inclusion/Representation of women",
                "questionAnswer":texts.get("participation/inclusion/representation of women","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Turnover rate for permanent employees and workers (Disclose trends for the past 3 years)",
                "questionAnswer":texts.get("turnover rate for permanent employees and workers (disclose trends for the past 3 years)","Not Applicable"),
              }
            ]
          },
          { "partRoman":"V",
            "categoryNo": "five",
            "subtitle": "Holding, Subsidiary and Associate Companies (including joint ventures)",
            "questions": [
              {
                "questionNo": "1",
                "question": "Names of holding/subsidiary/associate companies/joint ventures",
                "questionAnswer":texts.get("Names of holding/subsidiary/associate companies/joint ventures","Not Applicable"),
              }
            ]
          },
          {"partRoman":"VI",
            "categoryNo": "six",
            "subtitle": "CSR Details",
            "questions": [
            {
            "questionNo": "1",
            "question": "Whether CSR is applicable as per section 135 of Companies Act, 2013: (Yes/No)",
            "questionAnswer":texts.get("whether csr is applicable as per section 135 of companies act, 2013: (yes/no)","Not Applicable"),
            },
                        {
            "questionNo": "2",
            "question": "Turnover (in Rs.)",
            "questionAnswer":texts.get("turnover (in rs.)","Not Applicable"),
            },
            {
            "questionNo": "3",
            "question": "Net worth (in Rs.)",
            "questionAnswer":texts.get("net worth (in rs.)","Not Applicable"),
            },
            ]
          },
          { "partRoman":"VII",
            "categoryNo": "seven",
            "subtitle": "Transparency and Disclosures Compliances",
            "questions": [
              {
                "questionNo": "1",
                "question": "Complaints/Grievances on any of the principles (Principles 1 to 9) under the National Guidelines on Responsible Business Conduct:",
                "questionAnswer":texts.get("complaints/grievances on any of the principles (principles 1 to 9) under the national guidelines on responsible business conduct:","Not Applicable"),
              },
              {
            "questionNo": "2",
            "question": "Please indicate material responsible business conduct and sustainability issues pertaining to environmental and social matters that present a risk or an opportunity to your business, rationale for identifying the same, approach to adapt or mitigate the risk along-with its financial implications, as per the following format.",
            "questionAnswer":texts.get("please indicate material responsible business conduct and sustainability issues pertaining to environmental and social matters that present a risk or an opportunity to your business, rationale for identifying the same, approach to adapt or mitigate the risk along-with its financial implications, as per the following format.","Not Applicable"),
            },

            ]
          }
        ]},
    {
      "title": "MANAGEMENT AND PROCESS DISCLOSURES",
        "section": "SECTION B",
        "Categories": [
          { "partRoman":"I",
            "categoryNo": "one",
            "subtitle": "Policy and management processes",
            "questions": [
              {
                "questionNo": "1",
                "question": "Whether your entity’s policy/policies cover each principle and its core elements of the NGRBCs. (Yes/No)",
                "questionAnswer":texts.get("whether your entity’s policy/policies cover each principle and its core elements of the ngrbcs. (yes/no)","Not Applicable"),

              },
              {
                "questionNo": "2",
                "question": "Has the policy been approved by the Board? (Yes/No)",
                "questionAnswer":texts.get("has the policy been approved by the board? (yes/no)","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "Web Link of the Policies, if available.",                
                "questionAnswer":texts.get("web link of the policies, if available.","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Whether the entity has translated the policy into procedures. (Yes / No)",
                
                "questionAnswer":texts.get("whether the entity has translated the policy into procedures. (yes / no)","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Do the enlisted policies extend to your value chain partners? (Yes/No)",
                
                "questionAnswer":texts.get("do the enlisted policies extend to your value chain partners? (yes/no)","Not Applicable"),
              },
              {
                "questionNo": "6",
                "question": "Name of the national and international codes/ certifications/labels/ standards (e.g. Forest Stewardship Council, Fairtrade, Rainforest Alliance,Trustea) standards (e.g. SA 8000, OHSAS, ISO, BIS) adopted by your entity and mapped to each principle.",
                
                "questionAnswer":texts.get("name of the national and international codes/ certifications/labels/ standards (e.g. forest stewardship council, fairtrade, rainforest alliance,trustea) standards (e.g. sa 8000, ohsas, iso, bis) adopted by your entity and mapped to each principle.","Not Applicable"),
              },
              {
                "questionNo": "7",
                "question": "Specific commitments, goals and targets set by the entity with defined timelines, if any.",
                
                "questionAnswer":texts.get("specific commitments, goals and targets set by the entity with defined timelines, if any.","Not Applicable"),
              },
              {
                "questionNo": "8",
                "question": "Performance of the entity against the specific commitments, goals and targets along-with reasons in case the same are not met.",
                
                "questionAnswer":texts.get("performance of the entity against the specific commitments, goals and targets along-with reasons in case the same are not met.","Not Applicable"),
              },

            ]
          },
          {  "partRoman":"II",
            "categoryNo":"two",
            "subtitle": "Governance, leadership and oversight",
            "questions": [
              {
                "questionNo": "1",
                "question": "Statement by director responsible for the business responsibility report, highlighting ESG related challenges, targets and achievements (listed entity has flexibility regarding the placement of this disclosure)",
                
                "questionAnswer":texts.get("statement by director responsible for the business responsibility report, highlighting esg related challenges, targets and achievements (listed entity has flexibility regarding the placement of this disclosure)","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question": "Details of the highest authority responsible for implementation and oversight of the Business Responsibility policy (ies).",
                
                "questionAnswer":texts.get("details of the highest authority responsible for implementation and oversight of the business responsibility policy (ies).","Not Applicable"),

              },
              {
                "questionNo": "3",
                "question": "Does the entity have a specified Committee of the Board/ Director responsible for decision making on sustainability related issues? (Yes / No). If yes, provide details.",
                
                "questionAnswer":texts.get("does the entity have a specified committee of the board/ director responsible for decision making on sustainability related issues? (yes / no). if yes, provide details.","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Indicate whether review was undertaken by Director / Committee of the Board/ Any other Committee",
                "questionAnswer":texts.get("indicate whether review was undertaken by director / committee of the board/ any other committee","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Frequency(Annually/ Half yearly/ Quarterly/ Any other – please specify)",
                "questionAnswer":texts.get("frequency(annually/ half yearly/ quarterly/ any other – please specify)","Not Applicable"),
              },
              {
                "questionNo": "6",
                "question": "Has the entity carried out independent assessment/ evaluation of the working of its policies by an external agency? (Yes/No). If yes, provide name of the agency.",
                
                "questionAnswer":texts.get("has the entity carried out independent assessment/ evaluation of the working of its policies by an external agency? (yes/no). if yes, provide name of the agency.","Not Applicable"),
              },
              {
                "questionNo": "7",
                "question": "If answer to question (1) above is “No” i.e. not all Principles are covered by a policy, reasons to be stated, as below:",
                "questionAnswer":texts.get("if answer to question (1) above is “no” i.e. not all principles are covered by a policy, reasons to be stated, as below:","Not Applicable"),
              },
              {
                "questionNo": "8",
                "question": "Upstream (Suppliers & Logistics Partners)",                
                "questionAnswer":texts.get("upstream (suppliers & logistics partners)","Not Applicable"),
              },
              {
                "questionNo": "9",
                "question": "Downstream (Distributors & Customers)",                
                "questionAnswer":texts.get("downstream (distributors & customers)","Not Applicable"),
              },
              ]},]
        },
    {
        "title": "PRINCIPLE WISE PERFORMANCE DISCLOSURE ",
        "section": "SECTION C",
        "Categories": [
            {
            "partRoman":"PRINCIPLE I",
            "categoryNo": "one",
            "subtitle": "Businesses should conduct and govern themselves with integrity, and in a manner that is Ethical, Transparent and Accountable",
            "questions": [
              {
                "questionNo": "1",
                "question": "Percentage coverage by training and awareness programmes on any of the Principles during the financial year:",
                "questionAnswer":texts.get("percentage coverage by training and awareness programmes on any of the principles during the financial year:","Not Applicable"),
              },
              {
              "questionNo":"2",
              "question":"Details of fines / penalties /punishment/ award/ compounding fees/ settlement amount paid in proceedings (by the entity or by directors / KMPs) with regulators/ law enforcement agencies/ judicial institutions, in the financial year, in the following format (Note: the entity shall make disclosures on the basis of materiality as specified in Regulation 30 of SEBI (Listing Obligations and Disclosure Obligations) Regulations, 2015 and as disclosed on the entity’s website",
              "questionAnswer":texts.get("details of fines / penalties /punishment/ award/ compounding fees/ settlement amount paid in proceedings (by the entity or by directors / kmps) with regulators/ law enforcement agencies/ judicial institutions, in the financial year, in the following format (note: the entity shall make disclosures on the basis of materiality as specified in regulation 30 of sebi (listing obligations and disclosure obligations) regulations, 2015 and as disclosed on the entity’s website","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question":"Monetary",
                "questionAnswer":texts.get("monetary","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question":"Non-Monetary",
                "questionAnswer":texts.get("non-monetary","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Of the instances disclosed in Question 2 above, details of the Appeal/ Revision preferred in cases where monetary or non-monetary action has been appealed.",
                "questionAnswer":texts.get("of the instances disclosed in question 2 above, details of the appeal/ revision preferred in cases where monetary or non-monetary action has been appealed.","Not Applicable"),
              },
                {
                "questionNo": "6",
                "question": "Does the entity have an anti-corruption or anti-bribery policy? If yes, provide details in brief and if available, provide a web-link to the policy.",
                "questionAnswer":texts.get("does the entity have an anti-corruption or anti-bribery policy? if yes, provide details in brief and if available, provide a web-link to the policy.","Not Applicable"),
              },
                {
                "questionNo": "7",
                "question": "Number of Directors/KMPs/employees/workers against whom disciplinary action was taken by any law enforcement agency for the charges of bribery/ corruption",
                "questionAnswer":texts.get("number of directors/kmps/employees/workers against whom disciplinary action was taken by any law enforcement agency for the charges of bribery/ corruption","Not Applicable"),
              },
                {
                "questionNo": "8",
                "question": "Details of complaints with regard to conflict of interest",
                "questionAnswer":texts.get("details of complaints with regard to conflict of interest","Not Applicable"),
              },
                {
                "questionNo": "9",
                "question": "Provide details of any corrective action taken or underway on issues related to fines / penalties / action taken by regulators/ law enforcement agencies/ judicial institutions, on cases of corruption and conflicts of interest.",
                "questionAnswer":texts.get("provide details of any corrective action taken or underway on issues related to fines / penalties / action taken by regulators/ law enforcement agencies/ judicial institutions, on cases of corruption and conflicts of interest.","Not Applicable"),
              },
                {
                "questionNo": "10",
                "question": "Number of days of accounts payables ((Accounts payable *365) / Cost of goods/services procured) in the following format:",
                "questionAnswer":texts.get("number of days of accounts payables ((accounts payable *365) / cost of goods/services procured) in the following format:","Not Applicable"),
              },
                {
                "questionNo": "11",
                "question": "Provide details of concentration of purchases and sales with trading houses, dealers, and related parties along-with loans and advances & investments, with related parties, in the following format:",
                "questionAnswer":texts.get("provide details of concentration of purchases and sales with trading houses, dealers, and related parties along-with loans and advances & investments, with related parties, in the following format:","Not Applicable"),
              },
                {
                "questionNo": "12",
                "question": "Awareness programmes conducted for value chain partners on any of the Principles during the financial year:",
                "questionAnswer":texts.get("awareness programmes conducted for value chain partners on any of the principles during the financial year:","Not Applicable"),
              },
                {
                "questionNo": "13",
                "question": "Does the entity have processes in place to avoid/ manage conflict of interests involving members of the Board? (Yes/No) If Yes, provide details of the same.",
                "questionAnswer":texts.get("does the entity have processes in place to avoid/ manage conflict of interests involving members of the board? (yes/no) if yes, provide details of the same.","Not Applicable"),
              }]},
            {
            "partRoman":"PRINCIPLE II",
            "categoryNo": "two",
            "subtitle": "Businesses should provide goods and services in a manner that is sustainable and safe",
            "questions": [
              {
                "questionNo": "1",
                "question": "Percentage of R&D and capital expenditure (capex) investments in specific technologies to improve the environmental and social impacts of product and processes to total R&D and capex investments made by the entity, respectively.",
                "questionAnswer":texts.get("percentage of r&d and capital expenditure (capex) investments in specific technologies to improve the environmental and social impacts of product and processes to total r&d and capex investments made by the entity, respectively.","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question": "Does the entity have procedures in place for sustainable sourcing? (Yes/No)",
                "questionAnswer":texts.get("does the entity have procedures in place for sustainable sourcing? (yes/no)","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "If yes, what percentage of inputs were sourced sustainably?",
                "questionAnswer":texts.get("if there are any significant social or environmental concerns and/or risks arising from production or disposal of your products / services, as identified in the life cycle perspective / assessments (lca) or through any other means, briefly describe the same along-with action taken to mitigate the same.","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Describe the processes in place to safely reclaim your products for reusing, recycling and disposing at the end of life, for :",
                "questionAnswer":texts.get("describe the processes in place to safely reclaim your products for reusing, recycling and disposing at the end of life, for :","Not Applicable"),
              },

              {
                "questionNo": "5",
                "question": "Whether Extended Producer Responsibility (EPR) is applicable to the entity’s activities (Yes / No). If yes, whether the waste collection plan is in line with the Extended Producer Responsibility (EPR) plan submitted to Pollution Control Boards? If not, provide steps taken to address the same.",
                "questionAnswer":texts.get("whether extended producer responsibility (epr) is applicable to the entity’s activities (yes / no). if yes, whether the waste collection plan is in line with the extended producer responsibility (epr) plan submitted to pollution control boards? if not, provide steps taken to address the same.","Not Applicable"),
              },

              {
                "questionNo": "6",
                "question": "Has the entity conducted Life Cycle Perspective / Assessments (LCA) for any of its products (for manufacturing industry) or for its services (for service industry)? If yes, provide details in the following format?",
                "questionAnswer":texts.get("has the entity conducted life cycle perspective / assessments (lca) for any of its products (for manufacturing industry) or for its services (for service industry)? if yes, provide details in the following format?","Not Applicable"),
              },

              {
                "questionNo": "7",
                "question": "If there are any significant social or environmental concerns and/or risks arising from production or disposal of your products / services, as identified in the Life Cycle Perspective / Assessments (LCA) or through any other means, briefly describe the same along-with action taken to mitigate the same.",
                "questionAnswer":texts.get("if there are any significant social or environmental concerns and/or risks arising from production or disposal of your products / services, as identified in the life cycle perspective / assessments (lca) or through any other means, briefly describe the same along-with action taken to mitigate the same.","Not Applicable"),
              },

              {
                "questionNo": "8",
                "question": "Percentage of recycled or reused input material to total material (by value) used in production (for manufacturing industry) or providing services (for service industry)",
                "questionAnswer":texts.get("percentage of recycled or reused input material to total material (by value) used in production (for manufacturing industry) or providing services (for service industry)","Not Applicable"),
              },

              {
                "questionNo": "9",
                "question": "Of the products and packaging reclaimed at end of life of products, amount (in metric tonnes) reused, recycled, and safely disposed, as per the following format:",
                "questionAnswer":texts.get("of the products and packaging reclaimed at end of life of products, amount (in metric tonnes) reused, recycled, and safely disposed, as per the following format:","Not Applicable"),
              },
              {
                "questionNo": "10",
                "question": "Reclaimed products and their packaging materials (as percentage of products sold) for each product category.",
                "questionAnswer":texts.get("reclaimed products and their packaging materials (as percentage of products sold) for each product category.","Not Applicable"),
              },]
            },
            {
            "partRoman":"PRINCIPLE III",
            "categoryNo": "three",
            "subtitle": "Businesses should respect and promote the well-being of all employees, including those in their value chains",
            "questions": [
              {
                "questionNo": "1",
                "question":"Details of measures for the well-being of employees:",
                "questionAnswer":texts.get("details of measures for the well-being of employees:","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question":"Details of measures for the well-being of workers:",
                "questionAnswer":texts.get("details of measures for the well-being of workers:","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "Details of retirement benefits, for Current and Previous FY",
                "questionAnswer":texts.get("details of retirement benefits, for current and previous fy","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Accessibility of workplaces",
                "questionAnswer":texts.get("accessibility of workplaces","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Does the entity have an equal opportunity policy as per the Rights of Persons with Disabilities Act, 2016? If so, provide a web-link to the policy.",
                "questionAnswer":texts.get("does the entity have an equal opportunity policy as per the rights of persons with disabilities act, 2016? if so, provide a web-link to the policy.","Not Applicable"),
              },
              {
                "questionNo": "6",
                "question": "Return to work and Retention rates of permanent employees and workers that took parental leave.",
                "questionAnswer":texts.get("return to work and retention rates of permanent employees and workers that took parental leave.","Not Applicable"),
              },
              {
                "questionNo": "7",
                "question": "Is there a mechanism available to receive and redress grievances for the following categories of employees and worker? If yes, give details of the mechanism in brief",
                "questionAnswer":texts.get("is there a mechanism available to receive and redress grievances for the following categories of employees and worker? if yes, give details of the mechanism in brief","Not Applicable"),
              },
              {
                "questionNo": "8",
                "question": "Membership of employees and worker in association(s) or Unions recognised by the listed entity:",
                "questionAnswer":texts.get("membership of employees and worker in association(s) or unions recognised by the listed entity:","Not Applicable"),
              },
              {
                "questionNo": "9",
                "question": "Details of training given to employees and workers:",
                "questionAnswer":texts.get("details of training given to employees and workers:","Not Applicable"),
              },
              {
                "questionNo": "10",
                "question": "Details of performance and career development reviews of employees and workers:",
                "questionAnswer":texts.get("details of performance and career development reviews of employees and workers:","Not Applicable"),
              },
              {
                "questionNo": "11",
                "question":"Whether an occupational health and safety management system has been implemented by the entity? (Yes/ No). If yes, the coverage such system?",
                "questionAnswer":texts.get("whether an occupational health and safety management system has been implemented by the entity? (yes/ no). if yes, the coverage such system?","Not Applicable"),
              },
              {
                "questionNo": "12",
                "question":"What are the processes used to identify work-related hazards and assess risks on a routine and non-routine basis by the entity?",
                "questionAnswer":texts.get("what are the processes used to identify work-related hazards and assess risks on a routine and non-routine basis by the entity?","Not Applicable"),
              },
              {
                "questionNo": "13",
                "question":"Whether you have processes for workers to report the work related hazards and to remove themselves from such risks. (Y/N)",
                "questionAnswer":texts.get("whether you have processes for workers to report the work related hazards and to remove themselves from such risks. (y/n)","Not Applicable"),
              },
              {
                "questionNo": "14",
                "question":"Do the employees/ worker of the entity have access to non-occupational medical and healthcare services? (Yes/ No)",
                "questionAnswer":texts.get("do the employees/ worker of the entity have access to non-occupational medical and healthcare services? (yes/ no)","Not Applicable"),
              },
              {
                "questionNo": "15",
                "question": "Details of safety related incidents, in the following format:",
                "questionAnswer":texts.get("details of safety related incidents, in the following format:","Not Applicable"),
              },
              {
                "questionNo": "16",
                "question": "Describe the measures taken by the entity to ensure a safe and healthy work place",
                "questionAnswer":texts.get("describe the measures taken by the entity to ensure a safe and healthy work place","Not Applicable"),
              },
              {
                "questionNo": "17",
                "question": "Number of Complaints on the following made by employees and workers:",
                "questionAnswer":texts.get("number of complaints on the following made by employees and workers:","Not Applicable"),
              },
              {
                "questionNo": "18",
                "question": "Assessments for the year:",
                "questionAnswer":texts.get("assessments for the year:","Not Applicable"),
              },
              {
                "questionNo": "19",
                "question": "Provide details of any corrective action taken or underway to address safety-related incidents (if any) and on significant risks / concerns arising from assessments of health & safety practices and working conditions.",
                "questionAnswer":texts.get("provide details of any corrective action taken or underway to address safety-related incidents (if any) and on significant risks / concerns arising from assessments of health & safety practices and working conditions.","Not Applicable"),
              },
                {
                "questionNo": "20",
                "question": "Does the entity extend any life insurance or any compensatory package in the event of death of (A) Employees (Y/N)",
                "questionAnswer":texts.get("does the entity extend any life insurance or any compensatory package in the event of death of (a) employees (y/n)","Not Applicable"),
              },
              {
                "questionNo": "21",
                "question": "Provide the measures undertaken by the entity to ensure that statutory dues have been deducted and deposited by the value chain partners",
                "questionAnswer":texts.get("provide the measures undertaken by the entity to ensure that statutory dues have been deducted and deposited by the value chain partners","Not Applicable"),
              },
              {
                "questionNo": "22",
                "question": "Provide the number of employees / workers having suffered high consequence work-related injury / ill-health / fatalities (as reported in Q11 of Essential Indicators above), who have been are rehabilitated and placed in suitable employment or whose family members have been placed in suitable employment:",
                "questionAnswer":texts.get("provide the number of employees / workers having suffered high consequence work-related injury / ill-health / fatalities (as reported in q11 of essential indicators above), who have been are rehabilitated and placed in suitable employment or whose family members have been placed in suitable employment:","Not Applicable"),
              },
              {
                "questionNo": "23",
                "question": "Does the entity provide transition assistance programs to facilitate continued employability and the management of career endings resulting from retirement or termination of employment? (Yes/ No)",
                "questionAnswer":texts.get("does the entity provide transition assistance programs to facilitate continued employability and the management of career endings resulting from retirement or termination of employment? (yes/ no)","Not Applicable"),
              },
              {
                "questionNo": "24",
                "question": "Details on assessment of value chain partners:",
                "questionAnswer":texts.get("details on assessment of value chain partners:","Not Applicable"),
              },
              {
                "questionNo": "25",
                "question": "Provide details of any corrective actions taken or underway to address significant risks / concerns arising from assessments of health and safety practices and working conditions of value chain partners.",
                "questionAnswer":texts.get("provide details of any corrective actions taken or underway to address significant risks / concerns arising from assessments of health and safety practices and working conditions of value chain partners.","Not Applicable"),
              },]
            },
            {
            "partRoman":"PRINCIPLE IV",
            "categoryNo": "four",
            "subtitle": "Businesses should respect the interests of and be responsive to all its stakeholders",
            "questions": [
              {
                "questionNo": "1",
                "question": "Describe the processes for identifying key stakeholder groups of the entity.",
                "questionAnswer":texts.get("describe the processes for identifying key stakeholder groups of the entity.","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question": "List stakeholder groups identified as key for your entity and the frequency of engagement with each stakeholder group.",
                "questionAnswer":texts.get("list stakeholder groups identified as key for your entity and the frequency of engagement with each stakeholder group.","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "Provide the processes for consultation between stakeholders and the Board on economic, environmental, and social topics or if consultation is delegated, how is feedback from such consultations provided to the Board.",
                "questionAnswer":texts.get("provide the processes for consultation between stakeholders and the board on economic, environmental, and social topics or if consultation is delegated, how is feedback from such consultations provided to the board.","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Whether stakeholder consultation is used to support the identification and management of environmental, and social topics (Yes / No). If so, provide details of instances as to how the inputs received from stakeholders on these topics were incorporated into policies and activities of the entity",
                "questionAnswer":texts.get("whether stakeholder consultation is used to support the identification and management of environmental, and social topics (yes / no). if so, provide details of instances as to how the inputs received from stakeholders on these topics were incorporated into policies and activities of the entity","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Provide details of instances of engagement with, and actions taken to, address the concerns of vulnerable/ marginalized stakeholder groups.",
                "questionAnswer":texts.get("provide details of instances of engagement with, and actions taken to, address the concerns of vulnerable/ marginalized stakeholder groups.","Not Applicable"),
              },]
            },
            {
            "partRoman":"PRINCIPLE V",
            "categoryNo": "five",
            "subtitle": " Businesses should respect and promote human rights",
            "questions": [
              {
                "questionNo": "1",
                "question": "Employees and workers who have been provided training on human rights issues and policy(ies) of the entity, in the following format:",
                "questionAnswer":texts.get("employees and workers who have been provided training on human rights issues and policy(ies) of the entity, in the following format:","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question": "Details of minimum wages paid to employees and workers, in the following format:",
                "questionAnswer":texts.get("details of minimum wages paid to employees and workers, in the following format:","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "Details of remuneration/salary/wages, in the following format:",
                "questionAnswer":texts.get("details of remuneration/salary/wages","Not Applicable"),
              },
              {
                "questionNo": "4",
                "question": "Do you have a focal point (Individual/ Committee) responsible for addressing human rights impacts or issues caused or contributed to by the business? (Yes/ No)",
                "questionAnswer":texts.get("do you have a focal point (individual/ committee) responsible for addressing human rights impacts or issues caused or contributed to by the business? (yes/ no)","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Describe the internal mechanisms in place to redress grievances related to human rights issues.",
                "questionAnswer":texts.get("","Not Applicable"),
              },
              {
                "questionNo": "6",
                "question": "Number of Complaints on the following made by employees and workers:",
                "questionAnswer":texts.get("number of complaints on the following made by employees and workers:","Not Applicable"),
              },
              {
                "questionNo": "7",
                "question":"Complaints filed under the Sexual Harassment of Women at Workplace (Prevention, Prohibition and Redressal) Act, 2013, in the following format:",
                "questionAnswer":texts.get("complaints filed under the sexual harassment of women at workplace (prevention, prohibition and redressal) act, 2013, in the following format:","Not Applicable"),
              },
              {
                "questionNo": "8",
                "question": "Mechanisms to prevent adverse consequences to the complainant in discrimination and harassment cases.",
                "questionAnswer":texts.get("mechanisms to prevent adverse consequences to the complainant in discrimination and harassment cases.","Not Applicable"),
              },
              {
                "questionNo": "9",
                "question": "Do human rights requirements form part of your business agreements and contracts? (Yes/ No)",
                "questionAnswer":texts.get("do human rights requirements form part of your business agreements and contracts? (yes/ no)","Not Applicable"),
              },
              {
                "questionNo": "10",
                "question": "Assessments for the year:",
                "questionAnswer":texts.get("assessments for the year:","Not Applicable"),
              },
              {
                "questionNo": "11",
                "question": "Provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at Question 10 above.",
                "questionAnswer":texts.get("provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at question 10 above.","Not Applicable"),
              },
              {
                "questionNo": "12",
                "question": "Details of a business process being modified / introduced as a result of addressing human rights grievances/complaints.",
                "questionAnswer":texts.get("details of a business process being modified / introduced as a result of addressing human rights grievances/complaints.","Not Applicable"),
              },
              {
                "questionNo": "13",
                "question": "Details of the scope and coverage of any Human rights due-diligence conducted",
                "questionAnswer":texts.get("details of the scope and coverage of any human rights due-diligence conducted","Not Applicable"),
              },
              {
                "questionNo": "14",
                "question": "Is the premise/office of the entity accessible to differently abled visitors, as per the requirements of the Rights of Persons with Disabilities Act, 2016?",
                "questionAnswer":texts.get("is the premise/office of the entity accessible to differently abled visitors, as per the requirements of the rights of persons with disabilities act, 2016?","Not Applicable"),
              },
              {
                "questionNo": "15",
                "question": "Details on assessment of value chain partners:",
                "questionAnswer":texts.get("details on assessment of value chain partners:","Not Applicable"),
              },
              {
                "questionNo": "16",
                "question": "Provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at Question 4 above.",
                "questionAnswer":texts.get("provide details of any corrective actions taken or underway to address significant risks / concerns arising from the assessments at question 4 above.","Not Applicable"),
              },]
            },
            {  
          "partRoman":"PRINCIPLE VI",
          "categoryNo": "six",
          "subtitle": "Businesses should respect and make efforts to protect and restore",
          "questions": [
            {
              "questionNo": "1",
              "question": "Details of total energy consumption (in Joules or multiples) and energy intensity, in the following format:",
              "questionAnswer":texts.get("details of total energy consumption (in joules or multiples) and energy intensity, in the following format:","Not Applicable"),
            },
            {
              "questionNo": "2",
              "question": "Does the entity have any sites / facilities identified as designated consumers (DCs) under the Performance, Achieve and Trade (PAT) Scheme of the Government of India? (Y/N) If yes, disclose whether targets set under the PAT scheme have been achieved. In case targets have not been achieved, provide the remedial action taken, if any.",
              "questionAnswer":texts.get("does the entity have any sites / facilities identified as designated consumers (dcs) under the performance, achieve and trade (pat) scheme of the government of india? (y/n) if yes, disclose whether targets set under the pat scheme have been achieved. in case targets have not been achieved, provide the remedial action taken, if any.","Not Applicable"),
            },
            {
              "questionNo": "3",
              "question": "Provide details of the following disclosures related to water, in the following format:",
              "questionAnswer":texts.get("provide details of the following disclosures related to water, in the following format:","Not Applicable"),
            },

            {
              "questionNo": "4",
              "question": "Provide the following details related to water discharged:",
              "questionAnswer":texts.get("provide the following details related to water discharged:","Not Applicable"),
            },

            {
              "questionNo": "5",
              "question": "Has the entity implemented a mechanism for Zero Liquid Discharge? If yes, provide details of its coverage and implementation",
              "questionAnswer":texts.get("has the entity implemented a mechanism for zero liquid discharge? if yes, provide details of its coverage and implementation","Not Applicable"),
            },

            {
              "questionNo": "6",
              "question": "Please provide details of air emissions (other than GHG emissions) by the entity, in the following format:",
              "questionAnswer":texts.get("please provide details of air emissions (other than ghg emissions) by the entity, in the following format:","Not Applicable"),
            },

            {
              "questionNo": "7",
              "question": "Provide details of greenhouse gas emissions (Scope 1 and Scope 2 emissions) & its intensity, in the following format:",
              "questionAnswer":texts.get("provide details of greenhouse gas emissions (scope 1 and scope 2 emissions) & its intensity, in the following format:","Not Applicable"),
            },

            {
              "questionNo": "8",
              "question": "Does the entity have any project related to reducing Green House Gas emission? If Yes, then provide details.",
              "questionAnswer":texts.get("does the entity have any project related to reducing green house gas emission? if yes, then provide details.","Not Applicable"),
            },

            {
              "questionNo": "9",
              "question": "Provide details related to waste management by the entity, in the following format:",
              "questionAnswer":texts.get("provide details related to waste management by the entity, in the following format:","Not Applicable"),
            },

            {
              "questionNo": "10",
              "question": "Briefly describe the waste management practices adopted in your establishments. Describe the strategy adopted by your company to reduce usage of hazardous and toxic chemicals in your products and processes and the practices adopted to manage such wastes.",
              "questionAnswer":texts.get("briefly describe the waste management practices adopted in your establishments. describe the strategy adopted by your company to reduce usage of hazardous and toxic chemicals in your products and processes and the practices adopted to manage such wastes.","Not Applicable"),
            },

            {
              "questionNo": "11",
              "question": "If the entity has operations/offices in/around ecologically sensitive areas (such as national parks, wildlife sanctuaries, biosphere reserves, wetlands, biodiversity hotspots, forests, coastal regulation zones etc.) where environmental approvals / clearances are required, please specify details in the following format:",
              "questionAnswer":texts.get("if the entity has operations/offices in/around ecologically sensitive areas (such as national parks, wildlife sanctuaries, biosphere reserves, wetlands, biodiversity hotspots, forests, coastal regulation zones etc.) where environmental approvals / clearances are required, please specify details in the following format:","Not Applicable"),
            },

            {
              "questionNo": "12",
              "question": "Details of environmental impact assessments of projects undertaken by the entity based on applicable laws, in the current financial year:",
              "questionAnswer":texts.get("details of environmental impact assessments of projects undertaken by the entity based on applicable laws, in the current financial year:","Not Applicable"),
            },

            {
              "questionNo": "13",
              "question": "Is the entity compliant with the applicable environmental law/ regulations/ guidelines in India; such as the Water (Prevention and Control of Pollution) Act, Air (Prevention and Control of Pollution) Act, Environment protection act and rules thereunder (Y/N). If not, provide details of all such non-compliances, in the following format:",
              "questionAnswer":texts.get("is the entity compliant with the applicable environmental law/ regulations/ guidelines in india; such as the water (prevention and control of pollution) act, air (prevention and control of pollution) act, environment protection act and rules thereunder (y/n). if not, provide details of all such non-compliances, in the following format:","Not Applicable"),
            },

            {
              "questionNo": "14",
              "question": "Water withdrawal, consumption and discharge in areas of water stress (in kilolitres):",
              "questionAnswer":texts.get("water withdrawal, consumption and discharge in areas of water stress (in kilolitres):","Not Applicable"),
            },
            

            {
              "questionNo": "15",
              "question": "Please provide details of total Scope 3 emissions & its intensity, in the following format",
              "questionAnswer":texts.get("please provide details of total scope 3 emissions & its intensity, in the following format","Not Applicable"),
            },

            {
              "questionNo": "16",
              "question": "With respect to the ecologically sensitive areas reported at Question 10 of Essential Indicators above, provide details of significant direct & indirect impact of the entity on biodiversity in such areas along-with prevention and remediation activities",
              "questionAnswer":texts.get("with respect to the ecologically sensitive areas reported at question 10 of essential indicators above, provide details of significant direct & indirect impact of the entity on biodiversity in such areas along-with prevention and remediation activities","Not Applicable"),
            },

            {
              "questionNo": "17",
              "question": "If the entity has undertaken any specific initiatives or used innovative technology or solutions to improve resource efficiency, or reduce impact due to emissions / effluent discharge / waste generated, please provide details of the same as well as outcome of such initiatives, as per the following format:",
              "questionAnswer":texts.get("if the entity has undertaken any specific initiatives or used innovative technology or solutions to improve resource efficiency, or reduce impact due to emissions / effluent discharge / waste generated, please provide details of the same as well as outcome of such initiatives, as per the following format:","Not Applicable"),
            },

            {
              "questionNo": "18",
              "question": "Does the entity have a business continuity and disaster management plan? Give details in 100 words/ web link.",
              "questionAnswer":texts.get("does the entity have a business continuity and disaster management plan? give details in 100 words/ web link.","Not Applicable"),
            },

            {
              "questionNo": "19",
              "question": "Disclose any significant adverse impact to the environment, arising from the value chain of the entity. What mitigation or adaptation measures have been taken by the entity in this regard.",
              "questionAnswer":texts.get("disclose any significant adverse impact to the environment, arising from the value chain of the entity. what mitigation or adaptation measures have been taken by the entity in this regard.","Not Applicable"),
            },
            {
              "questionNo": "20",
              "question": "Percentage of value chain partners (by value of business done with such partners) that were assessed for environmental impacts.",
              "questionAnswer":texts.get("percentage of value chain partners (by value of business done with such partners) that were assessed for environmental impacts.","Not Applicable"),
            },
            {
              "questionNo": "21",
              "question": "How many Green Credits have been generated or procured:",
              "questionAnswer":texts.get("how many green credits have been generated or procured:","Not Applicable"),
            },
            {
              "questionNo": "22",
              "question": "By the listed entity",
              "questionAnswer":texts.get("by the listed entity","Not Applicable"),
            },
            {
              "questionNo": "23",
              "question": "By the top ten (in terms of value of purchases and sales,respectively) value chain partners",
              "questionAnswer":texts.get("By the top ten (in terms of value of purchases and sales,respectively) value chain partners","Not Applicable"),
            },]
          },
            {
          "partRoman":"PRINCIPLE VII",
          "categoryNo": "seven",
          "subtitle": " Businesses, when engaging in influencing public and regulatory policy, should do so in a manner that is responsible and transparent",
          "questions": [
            {
              "questionNo": "1",
              "question": "Number of affiliations with trade and industry chambers/ associations.",
              "questionAnswer":texts.get("number of affiliations with trade and industry chambers/ associations.","Not Applicable")
            },
            {
              "questionNo": "2",
              "question": "List the top 10 trade and industry chambers/ associations (determined based on the total members of such body) the entity is a member of/ affiliated to, in the following format",
              "questionAnswer":texts.get("list the top 10 trade and industry chambers/ associations (determined based on the total members of such body) the entity is a member of/ affiliated to, in the following format","Not Applicable")
            },
            
            {
              "questionNo": "3",
              "question": "Provide details of corrective action taken or underway on any issues related to anti-competitive conduct by the entity, based on adverse orders from regulatory authorities.",
              "questionAnswer":texts.get("provide details of corrective action taken or underway on any issues related to anti-competitive conduct by the entity, based on adverse orders from regulatory authorities.","Not Applicable")
            },
            {
              "questionNo": "4",
              "question": "Details of public policy positions advocated by the entity:",
              "questionAnswer":texts.get("details of public policy positions advocated by the entity:","Not Applicable")
            },]
          },
            {
            "partRoman":"PRINCIPLE VIII",
            "categoryNo": "eight",
            "subtitle": "Businesses should promote inclusive growth and equitable development",
            "questions": [
              {
                "questionNo": "1",
                "question": "Details of Social Impact Assessments (SIA) of projects undertaken by the entity based on applicable laws, in the current financial year.",
              "questionAnswer":texts.get("details of social impact assessments (sia) of projects undertaken by the entity based on applicable laws, in the current financial year.","Not Applicable"),
              },
              {
                "questionNo": "2",
                "question": "Provide information on project(s) for which ongoing Rehabilitation and Resettlement (R&R) is being undertaken by your entity, in the following format",
              "questionAnswer":texts.get("provide information on project(s) for which ongoing rehabilitation and resettlement (r&r) is being undertaken by your entity, in the following format","Not Applicable"),
                },

              {
                "questionNo": "3",
                "question": "Describe the mechanisms to receive and redress grievances of the community.",
              "questionAnswer":texts.get("describe the mechanisms to receive and redress grievances of the community.","Not Applicable"),
                },

              {
                "questionNo": "4",
                "question": "Percentage of input material (inputs to total inputs by value) sourced from suppliers",
              "questionAnswer":texts.get("percentage of input material (inputs to total inputs by value) sourced from suppliers","Not Applicable"),
                },

              {
                "questionNo": "5",
                "question": "Job creation in smaller towns – Disclose wages paid to persons employed (including employees or workers employed on a permanent or non-permanent / on contract basis) in the following locations, as % of total wage cost",
              "questionAnswer":texts.get("job creation in smaller towns – disclose wages paid to persons employed (including employees or workers employed on a permanent or non-permanent / on contract basis) in the following locations, as % of total wage cost","Not Applicable"),
                },

              {
                "questionNo": "6",
                "question": "Provide details of actions taken to mitigate any negative social impacts identified in the Social Impact Assessments (Reference: Question 1 of Essential Indicators above):",
              "questionAnswer":texts.get("provide details of actions taken to mitigate any negative social impacts identified in the social impact assessments (reference: question 1 of essential indicators above):","Not Applicable"),
                },

              {
                "questionNo": "7",
                "question": "Provide the following information on CSR projects undertaken by your entity in designated aspirational districts as identified by government bodies",
              "questionAnswer":texts.get("provide the following information on csr projects undertaken by your entity in designated aspirational districts as identified by government bodies","Not Applicable"),
              },
              {
                "questionNo": "8",
                "question": "Do you have a preferential procurement policy where you give preference to purchase from suppliers comprising marginalized /vulnerable groups? (Yes/No)",
              "questionAnswer":texts.get("do you have a preferential procurement policy where you give preference to purchase from suppliers comprising marginalized /vulnerable groups? (yes/no)","Not Applicable"),
              },

              {
                "questionNo": "9",
                "question": "From which marginalized /vulnerable groups do you procure?",
              "questionAnswer":texts.get("from which marginalized /vulnerable groups do you procure?","Not Applicable"),
              },
              {
                "questionNo": "10",
                "question": "What percentage of total procurement (by value) does it constitute?",
              "questionAnswer":texts.get("what percentage of total procurement (by value) does it constitute?","Not Applicable"),
              },
              {
                "questionNo": "11",
                "question": "Details of the benefits derived and shared from the intellectual properties owned or acquired by your entity (in the current financial year), based on traditional knowledge:",
              "questionAnswer":texts.get("details of the benefits derived and shared from the intellectual properties owned or acquired by your entity (in the current financial year), based on traditional knowledge:","Not Applicable"),
              },

              {
                "questionNo": "12",
                "question": "Details of corrective actions taken or underway, based on any adverse order in intellectual property related disputes wherein usage of traditional knowledge is involved.",
              "questionAnswer":texts.get("details of corrective actions taken or underway, based on any adverse order in intellectual property related disputes wherein usage of traditional knowledge is involved.","Not Applicable"),
              },

              {
                "questionNo": "13",
                "question": "Details of beneficiaries of CSR Projects:",
              "questionAnswer":texts.get("details of beneficiaries of csr projects:","Not Applicable"),
              },]
            },
            {
            "partRoman":"PRINCIPLE IX",
            "categoryNo": "nine",
            "subtitle": "Businesses should engage with and provide value to their consumers in a responsible manner",
            "questions": [
              {
                "questionNo": "1",
                "question": "Describe the mechanisms in place to receive and respond to consumer complaints and feedback.",
              "questionAnswer":texts.get("describe the mechanisms in place to receive and respond to consumer complaints and feedback.","Not Applicable"),
                },
              {
                "questionNo": "2",
                "question": "Turnover of products and/ services as a percentage of turnover from all products/service that carry information about:",
              "questionAnswer":texts.get("turnover of products and/ services as a percentage of turnover from all products/service that carry information about:","Not Applicable"),
              },
              {
                "questionNo": "3",
                "question": "Number of consumer complaints in respect of the following:",
              "questionAnswer":texts.get("number of consumer complaints in respect of the following:","Not Applicable"),
              },

              {
                "questionNo": "4",
                "question": "Details of instances of product recalls on account of safety issues:",
              "questionAnswer":texts.get("details of instances of product recalls on account of safety issues:","Not Applicable"),
              },
              {
                "questionNo": "5",
                "question": "Does the entity have a framework/policy on cyber security and risks related to data privacy? (Yes/No). If available, provide weblink of the policy.",
              "questionAnswer":texts.get("does the entity have a framework/policy on cyber security and risks related to data privacy? (yes/no). if available, provide weblink of the policy.","Not Applicable"),
              },

              {
                "questionNo": "6",
                "question": "Provide details of any corrective actions taken or underway on issues relating to advertising, and delivery of essential services; cyber security and data privacy of customers; re-occurrence of instances of product recalls; penalty / action taken by regulatory authorities on safety of products / services.",
              "questionAnswer":texts.get("provide details of any corrective actions taken or underway on issues relating to advertising, and delivery of essential services; cyber security and data privacy of customers; re-occurrence of instances of product recalls; penalty / action taken by regulatory authorities on safety of products / services.","Not Applicable"),
              },
              {
                "questionNo": "7",
                "question": "Channels / platforms where information on products and services of the entity can be accessed (provide web link, if available).",
              "questionAnswer":texts.get("channels / platforms where information on products and services of the entity can be accessed (provide web link, if available).","Not Applicable"),
              },

              {
                "questionNo": "8",
                "question": "Steps taken to inform and educate consumers about safe and responsible usage of products and/or services.",
              "questionAnswer":texts.get("steps taken to inform and educate consumers about safe and responsible usage of products and/or services.","Not Applicable"),
              },

              {
                "questionNo": "9",
                "question": "Mechanisms in place to inform consumers of any risk of disruption/discontinuation of essential services.",
              "questionAnswer":texts.get("mechanisms in place to inform consumers of any risk of disruption/discontinuation of essential services.","Not Applicable"),
              },

              {
                "questionNo": "10",
                "question": "Does the entity display product information on the product over and above what is mandated as per local laws? (Yes/No/Not Applicable) If yes, provide details in brief. Did your entity carry out any survey with regard to consumer satisfaction relating to the major products / services of the entity, significant locations of operation of the entity or the entity as a whole? (Yes/No)",
              "questionAnswer":texts.get("does the entity display product information on the product over and above what is mandated as per local laws? (yes/no/not applicable) if yes, provide details in brief. did your entity carry out any survey with regard to consumer satisfaction relating to the major products / services of the entity, significant locations of operation of the entity or the entity as a whole? (yes/no)","Not Applicable"),
              },
              {
                "questionNo": "11",
                "question": "Number of instances of data breaches along-with impact",
              "questionAnswer":texts.get("number of instances of data breaches along-with impact","Not Applicable"),
              },
              {
                "questionNo": "12",
                "question": "Percentage of data breaches involving personally identifiable information of customers",
              "questionAnswer":texts.get("provide the following information relating to data breaches:","Not Applicable"),
              },]
            }
          ]}
    ]}
        
     
     
        
          # 1. Save section

        table_name="Undefined BRSR"
        if payload.brsrfilename:
          filename =payload.brsrfilename.replace(".pdf", "")
          print("filename", filename)
          table_name = filename
        else:
          count_of_undefined=db.query(TableRegistry).filter_by(table_name=table_name).count()
          dup_int=count_of_undefined+1
          table_name="Undefined BRSR"+str(dup_int)
          
        PDFTable = create_pdf_model(table_name)
        sectionfind=sectionfind
        print("THis is section ###$$",sectionfind)
        # Track table creation order
        Base.metadata.create_all(bind=db.get_bind(), tables=[TableRegistry.__table__])
        existing = db.query(TableRegistry).filter_by(table_name=table_name).first()
        if not existing:
            registry_entry = TableRegistry(table_name=table_name,section=sectionfind)
            db.add(registry_entry)
            db.commit()

        # Create dynamic table if not exists
        PDFTable.__table__.create(bind=db.get_bind(), checkfirst=True)

        # Insert records
        for section_data in data.get("sections", []):
            if not section_data.get("title") or not section_data.get("section"):
                continue

            section = section_data["section"]
            title = section_data["title"]

            for category in section_data.get("Categories", []):
                if not category.get("categoryNo") or not category.get("subtitle"):
                    continue

                categoryNo = category["categoryNo"]
                subtitle = category["subtitle"]

                for q in category.get("questions", []):
                    if not q.get("questionNo") or not q.get("question") or "questionAnswer" not in q:
                        continue

                    answer = q["questionAnswer"]
                    if isinstance(answer, list):
                        answer = json.dumps(answer)

                    record = PDFTable(
                        section=section,
                        title=title,
                        categoryNo=categoryNo,
                        subtitle=subtitle,
                        question_no=q["questionNo"],
                        question=q["question"],
                        answer=answer
                    )
                    db.add(record)

        db.commit()
        return table_name



@app.get("/report_list/", response_model=ReportListResponse)
async def get_all_table_datas(db: Session = Depends(get_db)):
    """
    Returns all dynamically created table names from table_registry.
    """
    tables = db.query(TableRegistry).order_by(TableRegistry.created_at.asc()).all()

    if tables:
        report_items = [
            ReportItem(
                name=table.table_name,
                created_date=table.created_at,
                period="FY 2024 -2026",
                progress=pdf.random_number(),
                status="Saved",
                section=table.section
            )
            for table in tables
        ]

        report_list_model = ReportListResponse(reports=report_items)
        return report_list_model
    else:        
      return None

    


@app.get("/edit_pdf_report_get/{raw_table_name}")
async def edit_pdf_report_get(raw_table_name:str,db:Session=Depends(get_db)):
  PDFTable=create_pdf_model(raw_table_name)
  if raw_table_name not in Base.metadata.tables:
    return {"message":f"pdf not exist !"}
  inspector=inspect(db.get_bind())
  if not inspector.has_table(raw_table_name):
    return {"message":f"Table '{raw_table_name}'dos not exist in the database!"}
  rows=db.query(PDFTable).all()
  table_entry = db.query(TableRegistry).filter(TableRegistry.table_name == raw_table_name).first()

  data=[]
  section=table_entry.section
  filename=table_entry.table_name
  for row in rows:
    row_dict=row.__dict__.copy()
    row_dict.pop("_sa_instance_state", None)
    data.append(row_dict)
    
  return {"section":section,"data":data,"filename":filename}




@app.put("/edit_pdf_report_put/")
async def edit_pdf_report_put(raw_dict: Dict, db: Session = Depends(get_db)):
    print("Received:", raw_dict)
    edited=False
    texts = raw_dict.get("texts", {})
    sectionfind = raw_dict.get("sectionfind")
    currentsection = raw_dict.get("currentsection")
    filename = raw_dict.get("indexName")

    if not filename:
        raise HTTPException(status_code=400, detail="indexName (filename) is required")

    PDFTable = create_pdf_model(filename)

    # Check if table is in SQLAlchemy Base
    if filename not in Base.metadata.tables:
        raise HTTPException(status_code=404, detail=f"Table model for '{filename}' does not exist.")

    # Check if table exists in actual DB
    inspector = inspect(db.get_bind())
    if not inspector.has_table(filename):
        raise HTTPException(status_code=404, detail=f"Table '{filename}' does not exist in the database.")

    # Get section info from registry
    table_entry = db.query(TableRegistry).filter(TableRegistry.table_name == filename and TableRegistry.section == sectionfind).first()
    if not table_entry:
      raise HTTPException(status_code=404, detail="Table metadata not found in registry.")
    edited=True
    section = table_entry.section
    print("Table name:", table_entry.table_name)
    print("Section:", section)
    print("Starting updates...")


    print(texts)
    # Update each matching record
    updated_count = 0
    for question_text, new_answer in texts.items():
        record = db.query(PDFTable).filter(PDFTable.question == question_text).first()
        if record:
           # Safely handle list or string
          if isinstance(new_answer, list):
              record.answer = json.dumps(new_answer)
          else:
              record.answer = new_answer
          
          updated_count += 1
    db.commit()
    
    return {"filename":filename,"edited":edited}



@app.get("/download_pdf_report/{raw_table_name}")
async def download_pdf_report(raw_table_name: str, db: Session = Depends(get_db)):

    # Step 1: Dynamically create model
    PDFTable = create_pdf_model(raw_table_name)

    # Step 2: Check if model is in metadata
    if raw_table_name not in Base.metadata.tables:
        return {"message": f"Model for table '{raw_table_name}' is not registered"}

    # Step 3: Check if table exists in the database
    inspector = inspect(db.get_bind())
    if not inspector.has_table(raw_table_name):
        return {"message": f"Table '{raw_table_name}' does not exist in the database"}

    # Step 4: Query all rows
    rows = db.query(PDFTable).all()

    # Step 5: Convert to list of dicts
    data = []
    for row in rows:
        row_dict = row.__dict__.copy()
        row_dict.pop("_sa_instance_state", None)
        data.append(row_dict)

    # Step 6: Create PDF (optional)
    print("pdf name:", raw_table_name)
    pdf_path=pdf.create_pdf(data, raw_table_name)
    
    return FileResponse(
    path=pdf_path,
    filename=os.path.basename(pdf_path),
    media_type="application/pdf",
    headers={"Content-Disposition": f"attachment; filename={os.path.basename(pdf_path)}"}
)



@app.delete("/delete_pdf_report/{raw_table_name}")
async def delete_pdf_report(raw_table_name: str, db: Session = Depends(get_db)):
    # Step 1: Check if the table is registered in table_registry
    registry_entry = db.query(TableRegistry).filter_by(table_name=raw_table_name).first()
    if not registry_entry:
        raise HTTPException(status_code=404, detail=f"Table '{raw_table_name}' not found in registry.")

    # Step 2: Check if the table exists in the actual database
    inspector = inspect(db.get_bind())
    if not inspector.has_table(raw_table_name):
        raise HTTPException(status_code=404, detail=f"Table '{raw_table_name}' does not exist in the database.")

    # Step 3: Dynamically create model
    PDFTable = create_pdf_model(raw_table_name)

    # Step 4: Drop the table
    try:
        PDFTable.__table__.drop(bind=db.get_bind())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete table '{raw_table_name}': {str(e)}")

    # Step 5: Remove the table entry from table_registry
    db.delete(registry_entry)
    db.commit()

    return raw_table_name




if __name__ == "__main__":
    uvicorn.run("test:app",host="0.0.0.0", port=1000, reload=False, log_level="debug")